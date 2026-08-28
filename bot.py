import os
import json
import re
import base64
import zipfile
import pytz
from io import BytesIO
from xml.sax.saxutils import escape as xml_escape
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_file
import firebase_admin
from firebase_admin import credentials, firestore
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
import requests
import anthropic
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

# --- NOMBRES DE MESES EN ESPAÑOL ---
MESES_ESP = {
    1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
    5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
    9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
}

# --- INICIALIZAR FIREBASE ---
firebase_json = os.environ.get("FIREBASE_SERVICE_ACCOUNT_JSON")
if firebase_json:
    cred_dict = json.loads(firebase_json)
    cred = credentials.Certificate(cred_dict)
    firebase_admin.initialize_app(cred)
    db = firestore.client()

# --- FUNCION PARA CALCULOS DE FECHA Y HORA DE EVENTOS ---
def parsear_fecha_hora(texto):
    tz = pytz.timezone('America/Argentina/Buenos_Aires')
    ahora = datetime.now(tz)
    fecha_evento = None
    txt = texto.lower()

    if 'pasado mañana' in txt:
        fecha_evento = ahora + timedelta(days=2)
    elif 'mañana' in txt:
        fecha_evento = ahora + timedelta(days=1)
    elif 'hoy' in txt:
        fecha_evento = ahora
    else:
        match_fecha = re.search(r'\b(\d{1,2})/(\d{1,2})(?:/(\d{2,4}))?\b', txt)
        if match_fecha:
            dia = int(match_fecha.group(1))
            mes = int(match_fecha.group(2))
            anio = int(match_fecha.group(3)) if match_fecha.group(3) else ahora.year
            if anio < 100:
                anio += 2000
            try:
                fecha_temp = tz.localize(datetime(anio, mes, dia))
                if not match_fecha.group(3) and fecha_temp.date() < ahora.date():
                    anio += 1
                fecha_evento = tz.localize(datetime(anio, mes, dia))
            except ValueError:
                fecha_evento = None
        else:
            dias_semana = {
                'lunes': 0, 'martes': 1, 'miercoles': 2, 'miércoles': 2,
                'jueves': 3, 'viernes': 4, 'sabado': 5, 'sábado': 5, 'domingo': 6
            }
            for dia_nombre, dia_num in dias_semana.items():
                if dia_nombre in txt:
                    dias_hasta = (dia_num - ahora.weekday() + 7) % 7
                    if dias_hasta == 0:
                        dias_hasta = 7
                    fecha_evento = ahora + timedelta(days=dias_hasta)
                    break

    match_hora = re.search(r'\b(\d{1,2}):(\d{2})\b', txt)
    hora_str = None
    inicio = None
    fin = None
    hora_default = False

    if match_hora and fecha_evento:
        horas = int(match_hora.group(1))
        minutos = int(match_hora.group(2))
        if 0 <= horas <= 23 and 0 <= minutos <= 59:
            inicio = fecha_evento.replace(hour=horas, minute=minutos, second=0, microsecond=0)
            fin = inicio + timedelta(hours=1)
            hora_str = match_hora.group(0)
    elif fecha_evento and not match_hora:
        # No se especificó hora -> se asume 08:30 por defecto
        inicio = fecha_evento.replace(hour=8, minute=30, second=0, microsecond=0)
        fin = inicio + timedelta(hours=1)
        hora_str = "08:30"
        hora_default = True

    return inicio, fin, hora_str, fecha_evento, hora_default

# --- CREAR EVENTO EN GOOGLE CALENDAR ---
def agregar_evento_calendar(resumen, inicio, fin):
    gcal_json = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON")
    if not gcal_json:
        return False
    
    info = json.loads(gcal_json)
    scopes = ['https://www.googleapis.com/auth/calendar']
    creds = Credentials.from_service_account_info(info, scopes=scopes)
    service = build('calendar', 'v3', credentials=creds)
    
    evento = {
        'summary': resumen,
        'start': {'dateTime': inicio.isoformat()},
        'end': {'dateTime': fin.isoformat()},
    }
    
    service.events().insert(calendarId='casihyasociados@gmail.com', body=evento).execute()
    return True

# --- FUNCION PARA INTERPRETAR MONTOS EN CUALQUIER FORMATO ---
# Entiende: "1000000" / "1,000,000" / "1.000.000" / "1 millon" / "1.5 millones" / "500 mil"
def extraer_monto(texto):
    txt = texto.lower()

    # 1) "X millon(es)" -> multiplica por 1.000.000
    m = re.search(r'\d+(?:[.,]\d+)?\s*(?:millones|mill[oó]n|mill)\b(?:\s*de\s*pesos)?', txt)
    if m:
        numero = re.search(r'\d+(?:[.,]\d+)?', m.group(0)).group(0)
        monto = float(numero.replace(',', '.')) * 1_000_000
        texto_sin = texto[:m.start()] + texto[m.end():]
        return monto, texto_sin

    # 2) "X mil" -> multiplica por 1.000
    m = re.search(r'\d+(?:[.,]\d+)?\s*mil\b(?:\s*de\s*pesos)?', txt)
    if m:
        numero = re.search(r'\d+(?:[.,]\d+)?', m.group(0)).group(0)
        monto = float(numero.replace(',', '.')) * 1_000
        texto_sin = texto[:m.start()] + texto[m.end():]
        return monto, texto_sin

    # 3) numero plano, con o sin separadores de miles/decimales
    m = re.search(r'\d[\d.,]*', txt)
    if not m:
        return None, texto
    token = m.group(0)
    sep_count = token.count('.') + token.count(',')
    if sep_count == 0:
        monto = float(token)
    elif sep_count > 1:
        # mas de un separador -> todos son de miles (1.000.000 / 1,000,000)
        monto = float(re.sub(r'[.,]', '', token))
    else:
        sep_char = '.' if '.' in token else ','
        entero, frac = token.split(sep_char)
        if len(frac) == 3:
            # separador de miles (1.000 / 1,000)
            monto = float(entero + frac)
        else:
            # separador decimal (1250.50)
            monto = float(entero + '.' + frac)
    texto_sin = texto[:m.start()] + texto[m.end():]
    return monto, texto_sin

# --- RESPONDER POR WHATSAPP ---
def responder_whatsapp(chat_id, texto):
    instance_id = os.environ.get("GREEN_API_INSTANCE_ID")
    token = os.environ.get("GREEN_API_TOKEN")
    url = f"https://api.green-api.com/waInstance{instance_id}/sendMessage/{token}"
    try:
        r = requests.post(url, json={"chatId": chat_id, "message": texto}, timeout=15)
        print("[GREEN-API sendMessage] status=%s body=%s" % (r.status_code, r.text[:300]))
        return r.status_code == 200
    except Exception as e:
        print("[GREEN-API sendMessage] ERROR:", str(e))
        return False

# --- AVISO DE CLIENTE NUEVO ARCHIVADO (grupo de WhatsApp) ---
@app.route('/notificar-nuevo-cliente', methods=['POST', 'OPTIONS'])
def notificar_nuevo_cliente():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    d = request.get_json(silent=True) or {}
    servicios = ', '.join(d.get('servicios') or []) or '—'
    telefono = d.get('telefono') or d.get('tel') or '—'
    msg = (
        "🆕 *Nuevo cliente*\n\n"
        "👤 *Nombre:* %s\n" % d.get('nombre', '—') +
        "🪪 *DNI:* %s\n" % d.get('dni', '—') +
        "🎂 *Edad:* %s\n" % d.get('edad', '—') +
        "💼 *Situación laboral:* %s\n" % d.get('situacionLaboral', '—') +
        "🏠 *Domicilio:* %s, %s\n" % (d.get('domicilio', '—'), d.get('ciudad', '—')) +
        "📞 *Teléfono:* %s\n" % telefono +
        "⚖️ *Servicios:* %s\n" % servicios
    )
    nota = (d.get('notaAtencion') or '').strip()
    if nota:
        msg += "\n📌 *Documentación / a tener en cuenta:*\n*%s*" % nota
    destino = (d.get('destino') or '').strip() or os.environ.get("GREEN_API_GROUP_ID")
    ok = responder_whatsapp(destino, msg)
    resp = jsonify({"ok": ok, "grupo_configurado": bool(destino)})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

# --- FORMULARIO PUBLICO DE NUEVO CLIENTE (sin login) ---
# Lo llena el cliente mismo desde un link (nuevo-cliente.html), antes de venir
# al estudio. Va server-side con el Admin SDK -- asi no hace falta abrir las
# reglas de Firestore para escritura publica, todo pasa por este endpoint que
# valida lo minimo y arma el mismo documento que carga la secretaria a mano.
@app.route('/nuevo-cliente-remoto', methods=['POST', 'OPTIONS'])
def nuevo_cliente_remoto():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    d = request.get_json(silent=True) or {}

    def error(msg, code=400):
        resp = jsonify({"status": "error", "detalle": msg})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, code

    # Campo trampa: invisible para una persona, pero los bots que completan
    # todos los inputs de un formulario suelen llenarlo igual.
    if (d.get('sitioWeb') or '').strip():
        return error('Solicitud inválida.')

    nombre = (d.get('nombre') or '').strip()
    apellido = (d.get('apellido') or '').strip()
    if not nombre or not apellido:
        return error('Falta el nombre o el apellido.')

    tz = pytz.timezone('America/Argentina/Buenos_Aires')
    hoy = datetime.now(tz).strftime('%Y-%m-%d')

    data = {
        'nombre': ('%s %s' % (nombre, apellido)).strip(),
        'apellido': apellido,
        'dni': (d.get('dni') or '').strip(),
        'cuil': (d.get('cuil') or '').strip(),
        'edad': (d.get('edad') or '').strip(),
        'estadoCivil': (d.get('estadoCivil') or '').strip(),
        'nacionalidad': (d.get('nacionalidad') or 'Argentina').strip(),
        'situacionLaboral': (d.get('situacionLaboral') or '').strip(),
        'fechaNac': (d.get('fechaNac') or '').strip(),
        'tel': (d.get('tel') or '').strip(),
        'email': (d.get('email') or '').strip(),
        'domicilio': (d.get('domicilio') or '').strip(),
        'ciudad': (d.get('ciudad') or '').strip(),
        'empleador': (d.get('empleador') or '').strip(),
        'antiguedad': (d.get('antiguedad') or '').strip(),
        'categoria': (d.get('categoria') or '').strip(),
        'fechaIngreso': (d.get('fechaIngreso') or '').strip(),
        'fechaEgreso': (d.get('fechaEgreso') or '').strip(),
        'jornada': (d.get('jornada') or '').strip(),
        'tipo': 'posible',
        'origenFormulario': True,
        'expte': '', 'etapa': 'Cita', 'fecha': hoy,
        'drive': '', 'driveFolderId': '',
        'enRevision': True,
        'creadoEn': firestore.SERVER_TIMESTAMP,
    }

    try:
        db.collection('clientes').add(data)
    except Exception as e:
        print("Error guardando cliente remoto:", str(e))
        return error('No se pudo guardar. Probá de nuevo en un momento.', 500)

    grupo = os.environ.get("GREEN_API_GROUP_ID")
    if grupo:
        responder_whatsapp(grupo, "🆕 *Nuevo cliente desde el formulario web*\n\n👤 %s\n📞 %s\n\nRevisar en Inicio → Clientes nuevos por revisar." % (data['nombre'], data['tel'] or '—'))

    resp = jsonify({"status": "ok"})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

# --- AVISO DE RECORDATORIO (grupo de WhatsApp) ---
@app.route('/notificar-recordatorio', methods=['POST', 'OPTIONS'])
def notificar_recordatorio():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    d = request.get_json(silent=True) or {}
    msg = "🔔 *Recordatorio*\n\n" + "%s\n" % d.get('desc', '—')
    if d.get('hora'):
        msg += "🕐 %shs\n" % d.get('hora')
    destino = (d.get('destino') or '').strip() or os.environ.get("GREEN_API_GROUP_ID")
    ok = responder_whatsapp(destino, msg)
    resp = jsonify({"ok": ok, "grupo_configurado": bool(destino)})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

# --- SEGUIMIENTO PERSONAL DE TAREAS PENDIENTES ---
# Ademas del aviso diario a TODO el grupo, una tarea puede pedir "recordarme
# en X dias": ese dia puntual (fechaRecordatorio) se avisa aparte, por privado,
# al numero personal de seguimiento -- no al grupo. Se manda una sola vez
# (recordatorioEnviado) aunque el cron pase de largo esa fecha.
WHATSAPP_SEGUIMIENTO_TAREAS = os.environ.get("WHATSAPP_SEGUIMIENTO_PERSONAL") or "5493515102323@c.us"

@app.route('/revisar-seguimientos-tareas', methods=['GET'])
def revisar_seguimientos_tareas():
    tz = pytz.timezone('America/Argentina/Buenos_Aires')
    hoy = datetime.now(tz).strftime('%Y-%m-%d')
    avisadas = 0
    try:
        docs = db.collection('tareas_pendientes').where('fechaRecordatorio', '<=', hoy).stream()
        for doc in docs:
            t = doc.to_dict()
            if t.get('recordatorioEnviado'):
                continue
            desc = (t.get('desc') or '').strip()
            if not desc:
                continue
            msg = "⏰ *Seguimiento de tarea*\n\nPediste que te avise para consultar el avance de esto:\n\n%s" % desc
            if responder_whatsapp(WHATSAPP_SEGUIMIENTO_TAREAS, msg):
                doc.reference.update({'recordatorioEnviado': True})
                avisadas += 1
    except Exception as e:
        print("Error revisando seguimientos de tareas:", str(e))
    resp = jsonify({"avisadas": avisadas})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

# --- DOCTOR IA (chat de consultas del estudio) ---
SYSTEM_PROMPT_DOCTOR_IA = (
    "Sos 'Dr IA', el asistente interno del estudio jurídico Casih & Asociados, "
    "especializado en derecho laboral. El estudio ejerce en Córdoba Capital, Argentina "
    "-- además de la LCT y la legislación nacional, tené siempre presente la normativa "
    "y los criterios procesales propios de la provincia de Córdoba y su fuero laboral "
    "cuando la consulta lo requiera (procedimiento, aranceles, competencia, criterios "
    "de los tribunales locales). Te consulta el personal del estudio (abogados y "
    "secretaría), no clientes externos.\n\n"
    "Rigor y fuentes:\n"
    "- Fundamentá tus respuestas en fuentes formales y vigentes: citá la norma y el "
    "artículo concreto (ej. 'art. 245 LCT', 'art. X del CPT de Córdoba') en vez de "
    "afirmaciones sueltas sin respaldo.\n"
    "- No afirmes nada de lo que no estés razonablemente seguro. Si dudás si una norma "
    "sigue vigente, si hubo una reforma reciente, o si el criterio varía según el "
    "juzgado/fuero, decilo explícitamente -- nunca inventes un artículo, un fallo, una "
    "cifra o una fecha para completar la respuesta.\n"
    "- Tu conocimiento tiene una fecha de corte. Si la consulta puede depender de una "
    "reforma legal o un fallo muy reciente, aclará que conviene verificarlo porque "
    "podrías no tener esa actualización.\n\n"
    "Estilo según el tipo de consulta:\n"
    "- Si te piden REDACTAR un escrito (telegrama, denuncia, demanda, etc.), usá el "
    "registro formal y las formalidades propias de ese tipo de documento.\n"
    "- Si te consultan una DUDA o quieren entender un tema, respondé de forma clara y "
    "explicativa, en lenguaje llano, y ejemplificá con un caso concreto cuando ayude a "
    "que se entienda mejor.\n"
    "- En temas con más de un criterio posible o zonas grises, decilo -- no elijas una "
    "sola postura como si fuera la única, salvo que la ley sea clara al respecto.\n\n"
    "- No tenés acceso a los datos de clientes ni expedientes del sistema. Si te "
    "preguntan algo puntual de un caso concreto (fechas, montos, números de expediente), "
    "aclará que no tenés esa información y que hay que revisarla en el sistema, salvo "
    "que te la hayan adjuntado directamente como imagen, PDF o Word en el mensaje -- en "
    "ese caso analizala.\n"
    "- Recordá siempre, aunque sea brevemente, que esto es una orientación de apoyo y "
    "no reemplaza la revisión final de un abogado antes de presentar algo o de asesorar "
    "a un cliente con eso."
)

@app.route('/doctor-ia', methods=['POST', 'OPTIONS'])
def doctor_ia():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    d = request.get_json(silent=True) or {}
    pregunta = (d.get('mensaje') or '').strip()
    adjunto = d.get('adjunto') or {}
    mime = (adjunto.get('mimeType') or '').strip()
    data_b64 = (adjunto.get('data') or '').strip()
    if not pregunta and not (mime and data_b64):
        resp = jsonify({"error": "Falta el mensaje."})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 400

    mensajes = []
    for h in (d.get('historial') or [])[-60:]:
        role = h.get('role')
        texto = (h.get('texto') or '').strip()
        if role in ('user', 'assistant') and texto:
            mensajes.append({"role": role, "content": texto})

    if mime and data_b64:
        bloque_archivo = None
        texto_docx = None
        if mime == 'application/pdf':
            bloque_archivo = {"type": "document", "source": {"type": "base64", "media_type": mime, "data": data_b64}}
        elif mime.startswith('image/'):
            bloque_archivo = {"type": "image", "source": {"type": "base64", "media_type": mime, "data": data_b64}}
        elif mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            try:
                doc = Document(BytesIO(base64.b64decode(data_b64)))
                partes = [p.text for p in doc.paragraphs if p.text.strip()]
                for tabla in doc.tables:
                    for fila in tabla.rows:
                        partes.append(' | '.join(c.text.strip() for c in fila.cells))
                texto_docx = '\n'.join(partes).strip() or '(el documento parece estar vacío)'
            except Exception as e:
                print("Error leyendo Word adjunto:", str(e))
                texto_docx = None

        if bloque_archivo:
            contenido = [bloque_archivo]
            if pregunta:
                contenido.append({"type": "text", "text": pregunta})
            mensajes.append({"role": "user", "content": contenido})
        elif texto_docx is not None:
            texto_final = "Contenido del archivo Word adjunto:\n\n%s" % texto_docx
            if pregunta:
                texto_final += "\n\n---\n\n%s" % pregunta
            mensajes.append({"role": "user", "content": texto_final})
        else:
            mensajes.append({"role": "user", "content": pregunta or "No se pudo leer el archivo adjunto."})
    else:
        mensajes.append({"role": "user", "content": pregunta})

    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        resp = jsonify({"error": "Falta configurar ANTHROPIC_API_KEY en el servidor."})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 500

    try:
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-sonnet-5",
            max_tokens=1536,
            system=SYSTEM_PROMPT_DOCTOR_IA,
            messages=mensajes
        )
        respuesta = "".join([b.text for b in response.content if b.type == "text"]).strip()
    except Exception as e:
        print("Error en Dr IA:", str(e))
        resp = jsonify({"error": "No se pudo obtener respuesta. Probá de nuevo."})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 500

    resp = jsonify({"respuesta": respuesta})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

# --- RECORDATORIOS DE HOY (para sumar a la agenda diaria) ---
@app.route('/recordatorios-del-dia', methods=['GET'])
def recordatorios_del_dia():
    tz = pytz.timezone('America/Argentina/Buenos_Aires')
    hoy = datetime.now(tz).strftime('%Y-%m-%d')
    items = []
    try:
        docs = db.collection('recordatorios').where('fecha', '==', hoy).stream()
        for doc in docs:
            r = doc.to_dict()
            titulo = (r.get('titulo') or '').strip()
            desc = (r.get('desc') or '').strip()
            cliente = (r.get('cliente') or '').strip()
            if titulo:
                texto = titulo
                if desc:
                    texto += ' — ' + desc
                if cliente:
                    texto += ' (Cliente: %s)' % cliente
            else:
                texto = desc
            if texto:
                items.append({'texto': texto, 'hora': r.get('hora') or ''})
    except Exception as e:
        print("Error obteniendo recordatorios del dia:", str(e))
    items.sort(key=lambda x: x['hora'] or '99:99')
    resp = jsonify(items)
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

# --- TAREAS PENDIENTES (para sumar a la agenda diaria) ---
# A diferencia de los recordatorios, no tienen fecha: existen mientras no se
# marquen como hechas (se borran desde el sistema al completarlas), asi que
# ninguna se filtra por dia -- se devuelven todas, para que se avisen todos
# los dias hasta que se hagan.
@app.route('/tareas-pendientes', methods=['GET'])
def tareas_pendientes():
    items = []
    tz = pytz.timezone('America/Argentina/Buenos_Aires')
    hoy = datetime.now(tz).date()
    try:
        docs = db.collection('tareas_pendientes').order_by('creadoEn').stream()
        for doc in docs:
            t = doc.to_dict()
            desc = (t.get('desc') or '').strip()
            if not desc:
                continue
            texto = desc
            fecha_plazo = (t.get('fechaPlazoIdeal') or '').strip()
            if fecha_plazo:
                try:
                    dias_restantes = (datetime.strptime(fecha_plazo, '%Y-%m-%d').date() - hoy).days
                    if dias_restantes < 0:
                        texto = '🚨 FUERA DE PLAZO (venció hace %d día/s) - %s' % (-dias_restantes, desc)
                    elif dias_restantes <= 2:
                        texto = '⚠️ Se acerca el plazo ideal (%s) - %s' % (
                            'hoy' if dias_restantes == 0 else 'quedan %d día/s' % dias_restantes, desc)
                except ValueError:
                    pass
            items.append({'texto': texto})
    except Exception as e:
        print("Error obteniendo tareas pendientes:", str(e))
    resp = jsonify(items)
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

# --- RUTA PRINCIPAL (WEBHOOK) ---
@app.route('/webhook', methods=['POST'])
def webhook():
    data = request.json
    print("[WEBHOOK recibido]", json.dumps(data)[:600] if data else "sin body/JSON")
    try:
        type_webhook = data.get('typeWebhook')
        if type_webhook == 'incomingMessageReceived':
            message_data = data.get('messageData', {})
            chat_id = data.get('senderData', {}).get('chatId')
            
            text_message = ""
            if 'textMessageData' in message_data:
                text_message = message_data['textMessageData'].get('textMessage', '').strip()
            elif 'extendedTextMessageData' in message_data:
                text_message = message_data['extendedTextMessageData'].get('text', '').strip()

            if not text_message:
                return jsonify({"status": "ignored"}), 200

            text_lower = text_message.lower()
            tz = pytz.timezone('America/Argentina/Buenos_Aires')
            ahora = datetime.now(tz)

            # Mismo formato que usa index.html: fecha "YYYY-MM-DD", mes de 2 dígitos,
            # año como string, y mesKey = "YYYY-MM" (es el campo que filtra Administración)
            fecha_hoy_str = ahora.strftime("%Y-%m-%d")
            mes_str = ahora.strftime("%m")
            anio_str = ahora.strftime("%Y")
            mes_key = f"{anio_str}-{mes_str}"

            # 0. COMANDO: AYUDA
            if text_lower.startswith('ayuda'):
                responder_whatsapp(
                    chat_id,
                    "🤖 *Comandos disponibles*\n\n"
                    "💰 *Gasto*\n`gasto MONTO CONCEPTO`\n_Ej: gasto 12000 librería_\n\n"
                    "💵 *Ingreso / Honorario*\n`ingreso MONTO CLIENTE`\n_Ej: ingreso 150000 Garcia_\n\n"
                    "📅 *Evento* (se crea en Google Calendar)\n`evento DÍA [HORA] TÍTULO`\n_Ej: evento miercoles 17:00 cita cliente Walter Figueroa_\n\n"
                    "⏰ *Recordatorio* (llega por WhatsApp a esta misma hora)\n`recordame DÍA [HORA] TEXTO`\n_Ej: recordame mañana 15:00 llamar a García_\n\n"
                    "📆 *Formas de indicar el día:* hoy, mañana, pasado mañana, un día de la semana (lunes, martes...), o una fecha exacta (15/08 o 15/08/2026).\n\n"
                    "🕐 Si no indicás la hora, se asume automáticamente las *08:30 hs*.\n\n"
                    "⚠️ El comando va siempre *al principio* del mensaje, y la hora en formato 24hs (15:00, no 3pm)."
                )

            # 1. COMANDO: GASTO
            elif text_lower.startswith('gasto'):
                monto, resto = extraer_monto(text_message)
                concepto = re.sub(r'gasto', '', resto, flags=re.IGNORECASE).strip()
                concepto = re.sub(r'\s{2,}', ' ', concepto).strip()

                if monto is None:
                    responder_whatsapp(chat_id, "⚠️ *Error al registrar gasto:*\nFalta indicar el *monto* numérico.\n\n💡 *Ejemplo:* `gasto 12000 librería`")
                elif not concepto:
                    responder_whatsapp(chat_id, "⚠️ *Error al registrar gasto:*\nFalta indicar el *concepto o detalle*.\n\n💡 *Ejemplo:* `gasto 12000 resma de hojas`")
                else:
                    doc_gasto = {
                        'fecha': fecha_hoy_str,
                        'concepto': concepto,
                        'cat': 'varios',            # clave usada por catLbl/catCls en index.html
                        'cargadoPor': 'Bot WhatsApp',
                        'monto': monto,
                        'mes': mes_str,              # "08" (2 dígitos, no "Agosto")
                        'anio': anio_str,            # "2026"
                        'mesKey': mes_key,           # "2026-08" -> esto es lo que filtra Administración
                        'creadoEn': firestore.SERVER_TIMESTAMP
                    }

                    db.collection('gastos').add(doc_gasto)

                    responder_whatsapp(chat_id, f"✅ *Gasto registrado correctamente*\n💰 *Monto:* ${monto:.2f}\n📝 *Concepto:* {concepto}\n📅 *Fecha:* {fecha_hoy_str}")

            # 2. COMANDO: INGRESO / HONORARIOS
            elif text_lower.startswith('ingreso') or text_lower.startswith('honorario'):
                monto, resto = extraer_monto(text_message)
                cliente = re.sub(r'ingreso|honorarios|honorario', '', resto, flags=re.IGNORECASE).strip()
                cliente = re.sub(r'\s{2,}', ' ', cliente).strip()

                if monto is None:
                    responder_whatsapp(chat_id, "⚠️ *Error al registrar ingreso:*\nFalta indicar el *monto* numérico.\n\n💡 *Ejemplo:* `ingreso 150000 Garcia`")
                else:
                    # El cliente es opcional (igual que en el formulario web) para no trabar la carga
                    concepto_txt = f"Cobro honorarios {cliente}" if cliente else "Cobro de honorarios"

                    doc_ingreso = {
                        'fecha': fecha_hoy_str,
                        'concepto': concepto_txt,
                        'cliente': cliente,          # puede quedar como ''
                        'cat': 'honorarios',          # clave usada por catLbl/catCls en index.html
                        'monto': monto,
                        'mes': mes_str,               # "08"
                        'anio': anio_str,              # "2026"
                        'mesKey': mes_key,            # "2026-08"
                        'creadoEn': firestore.SERVER_TIMESTAMP
                    }

                    db.collection('ingresos').add(doc_ingreso)

                    linea_cliente = f"\n👤 *Cliente:* {cliente}" if cliente else ""
                    responder_whatsapp(chat_id, f"💵 *Ingreso registrado correctamente*\n💰 *Monto:* ${monto:.2f}{linea_cliente}\n📝 *Concepto:* {concepto_txt}\n📅 *Fecha:* {fecha_hoy_str}")

            # 3. COMANDO: EVENTO
            elif text_lower.startswith('evento'):
                inicio, fin, hora_str, fecha_detectada, hora_default = parsear_fecha_hora(text_message)
                
                resumen = re.sub(
                    r'evento|mañana|pasado mañana|hoy|lunes|martes|miercoles|miércoles|jueves|viernes|sabado|sábado|domingo|\b\d{1,2}:\d{2}\b|\b\d{1,2}/\d{1,2}(?:/\d{2,4})?\b', 
                    '', 
                    text_message, 
                    flags=re.IGNORECASE
                ).strip()

                if not fecha_detectada:
                    responder_whatsapp(
                        chat_id, 
                        "⚠️ *Error al crear evento:*\nNo se detectó el *día o fecha*.\n\n"
                        "💡 *Ejemplo:* `evento miercoles 17:00 cita cliente Walter Figueroa`"
                    )
                elif not resumen:
                    responder_whatsapp(
                        chat_id, 
                        "⚠️ *Error al crear evento:*\nFalta indicar el *motivo o título*.\n\n"
                        "💡 *Ejemplo:* `evento mañana 16:30 Reunion with Client`"
                    )
                else:
                    try:
                        if agregar_evento_calendar(resumen, inicio, fin):
                            fecha_formateada = inicio.strftime("%d/%m/%Y")
                            nota_default = " _(hora no indicada, se asumió por defecto)_" if hora_default else ""
                            responder_whatsapp(
                                chat_id, 
                                f"📅 *Evento agendado en Google Calendar:*\n"
                                f"📆 *Fecha:* {fecha_formateada}\n"
                                f"⏰ *Hora:* {hora_str} hs{nota_default}\n"
                                f"📌 *Título:* {resumen}"
                            )
                        else:
                            responder_whatsapp(chat_id, "❌ *Error:* No se pudo conectar con Google Calendar.")
                    except Exception as err:
                        responder_whatsapp(chat_id, f"❌ *Error al crear evento:* {str(err)}")

            # 4. COMANDO: RECORDATORIO
            elif text_lower.startswith('recordame') or text_lower.startswith('recordá') or text_lower.startswith('recorda'):
                inicio, fin, hora_str, fecha_detectada, hora_default = parsear_fecha_hora(text_message)

                texto_recordatorio = re.sub(
                    r'recordame|recorda|recordá|mañana|pasado mañana|hoy|lunes|martes|miercoles|miércoles|jueves|viernes|sabado|sábado|domingo|\b\d{1,2}:\d{2}\b|\b\d{1,2}/\d{1,2}(?:/\d{2,4})?\b',
                    '',
                    text_message,
                    flags=re.IGNORECASE
                ).strip()

                if not fecha_detectada:
                    responder_whatsapp(
                        chat_id,
                        "⚠️ *Error al crear recordatorio:*\nNo se detectó el *día o fecha*.\n\n"
                        "💡 *Ejemplo:* `recordame mañana 15:00 llamar a García`"
                    )
                elif not texto_recordatorio:
                    responder_whatsapp(
                        chat_id,
                        "⚠️ *Error al crear recordatorio:*\nFalta indicar *qué* recordar.\n\n"
                        "💡 *Ejemplo:* `recordame mañana 15:00 llamar a García`"
                    )
                else:
                    doc_recordatorio = {
                        'chatId': chat_id,           # a este mismo chat (privado o grupo) se le responde
                        'fecha': inicio.strftime("%Y-%m-%d"),
                        'hora': inicio.strftime("%H:%M"),
                        'texto': texto_recordatorio,
                        'enviado': False,
                        'creadoEn': firestore.SERVER_TIMESTAMP
                    }
                    db.collection('recordatorios_bot').add(doc_recordatorio)

                    fecha_formateada = inicio.strftime("%d/%m/%Y")
                    nota_default = " _(hora no indicada, se asumió por defecto)_" if hora_default else ""
                    responder_whatsapp(
                        chat_id,
                        f"⏰ *Recordatorio guardado*\n"
                        f"📆 *Fecha:* {fecha_formateada}\n"
                        f"🕐 *Hora:* {hora_str} hs{nota_default}\n"
                        f"📝 *Texto:* {texto_recordatorio}"
                    )

    except Exception as e:
        print("Error procesando Webhook:", str(e))
        
    return jsonify({"status": "success"}), 200

# --- RUTA PARA REVISAR Y ENVIAR RECORDATORIOS PENDIENTES ---
# La llama cron-job.org cada 5 minutos (GET simple). De paso, ese mismo
# ping mantiene despierto el servicio en Render, que se duerme tras 15
# minutos sin actividad en el plan gratuito.
@app.route('/revisar-recordatorios', methods=['GET'])
def revisar_recordatorios():
    tz = pytz.timezone('America/Argentina/Buenos_Aires')
    ahora = datetime.now(tz)
    enviados = 0
    errores = 0

    try:
        pendientes = db.collection('recordatorios_bot').where('enviado', '==', False).stream()
        for doc in pendientes:
            data = doc.to_dict()
            try:
                fecha_hora_str = f"{data['fecha']} {data['hora']}"
                programado = tz.localize(datetime.strptime(fecha_hora_str, "%Y-%m-%d %H:%M"))

                if ahora >= programado:
                    ok = responder_whatsapp(
                        data['chatId'],
                        f"⏰ *Recordatorio:*\n{data['texto']}"
                    )
                    if ok:
                        doc.reference.update({'enviado': True, 'enviadoEn': firestore.SERVER_TIMESTAMP})
                        enviados += 1
                    else:
                        # No se pudo confirmar la entrega (ej. cupo de Green API agotado):
                        # se deja pendiente para reintentar en la proxima corrida del cron.
                        print("Recordatorio no confirmado, se reintenta:", doc.id)
                        errores += 1
            except Exception as e_inner:
                print("Error procesando recordatorio", doc.id, str(e_inner))
                errores += 1
    except Exception as e:
        print("Error revisando recordatorios:", str(e))
        return jsonify({"status": "error", "detalle": str(e)}), 500

    return jsonify({"status": "ok", "enviados": enviados, "errores": errores, "hora_chequeo": ahora.isoformat()}), 200

# --- SEGUIMIENTO CORREO ARGENTINO ---
# El formulario publico de seguimiento no pide captcha: es un POST simple que
# devuelve un fragmento HTML con la tabla de movimientos (mas nuevo primero).
CA_URL = "https://www.correoargentino.com.ar/sites/all/modules/custom/ca_forms/api/wsFacade.php"

def consultar_correo_ca(prefijo, numero):
    headers = {
        'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
        'X-Requested-With': 'XMLHttpRequest',
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36',
        'Referer': 'https://www.correoargentino.com.ar/formularios/ondnc'
    }
    payload = {'action': 'ondnc', 'id': numero, 'producto': prefijo, 'pais': 'AR'}
    resp = requests.post(CA_URL, data=payload, headers=headers, timeout=30)
    resp.raise_for_status()
    html = resp.content.decode('utf-8-sig', errors='replace')

    if 'No se encontraron resultados' in html:
        return []

    filas = re.findall(
        r'data-title="Fecha:">(.*?)</td>\s*'
        r'<td data-title="Planta:">(.*?)</td>\s*'
        r'<td data-title="Historia:">(.*?)</td>\s*'
        r'<td data-title="Estado:">(.*?)</td>',
        html, re.S
    )
    return [
        {'fecha': f.strip(), 'planta': p.strip(), 'historia': h.strip(), 'estado': e.strip()}
        for f, p, h, e in filas
    ]

def procesar_pieza_correo(doc, data, grupo):
    """Consulta una pieza, actualiza Firestore y devuelve que paso.
    Resultado: 'error' | 'sin_resultados' | 'primera_vez' | 'sin_cambios' | 'con_novedad'
    """
    prefijo = data.get('caPrefijo')
    numero = data.get('caNumero')

    try:
        movs = consultar_correo_ca(prefijo, numero)
    except Exception as e_req:
        print("Error consultando", prefijo, numero, str(e_req))
        doc.reference.update({'caError': 'No se pudo consultar', 'caUltimaConsulta': firestore.SERVER_TIMESTAMP})
        return 'error'

    if not movs:
        doc.reference.update({'caError': 'Sin resultados en Correo Argentino', 'caUltimaConsulta': firestore.SERVER_TIMESTAMP})
        return 'sin_resultados'

    ultimo = movs[0]
    # Identifica el movimiento de arriba de la tabla (no solo cuantos hay en
    # total). El Correo a veces carga un movimiento anterior en el tiempo
    # *despues* de uno posterior -- por ej. "LLEGADA AL CENTRO DE
    # PROCESAMIENTO" aparece en el sistema recien despues de que ya figuraba
    # "INTENTO DE ENTREGA" con la misma fecha -- y ahi la cantidad total de
    # movimientos aumenta sin que el de arriba haya cambiado. Comparar solo
    # la cantidad hacia que se reenviara el mismo aviso de nuevo en esos
    # casos; comparando la identidad del ultimo movimiento no.
    clave_actual = '%s|%s|%s|%s' % (ultimo['fecha'], ultimo['planta'], ultimo['historia'], ultimo['estado'])
    clave_previa = data.get('caUltimoMovKey') or ''

    update = {
        'caEstado': ultimo['estado'],
        'caHistoria': ultimo['historia'],
        'caPlanta': ultimo['planta'],
        'caFechaMov': ultimo['fecha'],
        'caMovimientos': len(movs),
        'caUltimoMovKey': clave_actual,
        'caError': '',
        'caUltimaConsulta': firestore.SERVER_TIMESTAMP
    }

    # Una vez entregada no se consulta mas
    if ultimo['estado'].upper() == 'ENTREGADO':
        update['caFinalizado'] = True

    # La primera consulta solo deja registrado el estado actual: avisar de
    # piezas recien cargadas seria ruido, no novedad.
    primera_vez = not data.get('caConsultado')
    # Piezas que ya se venian trackeando antes de que existiera caUltimoMovKey
    # no tienen con que comparar todavia: se toma esta corrida como el punto
    # de partida (sin avisar) en vez de asumir que todo es "novedad".
    sin_clave_previa = data.get('caConsultado') and not data.get('caUltimoMovKey')
    update['caConsultado'] = True

    doc.reference.update(update)

    if primera_vez or sin_clave_previa:
        return 'primera_vez'

    if clave_actual == clave_previa:
        return 'sin_cambios'

    estado = ultimo['estado'].upper()
    historia = ultimo['historia'].upper()
    # Solo interesan los intentos de entrega, el rechazo y la entrega final.
    # El resto (llegada a planta, clasificacion, en poder del cartero, etc.)
    # son movimientos internos del Correo que no requieren que nadie del
    # estudio haga nada, asi que no generan WhatsApp -- pero igual quedan
    # registrados arriba (caUltimoMovKey ya se actualizo) para no reenviar
    # un aviso viejo cuando despues aparezca uno que si importe.
    es_relevante = estado in ('RECHAZADO', 'ENTREGADO') or 'INTENTO DE ENTREGA' in historia

    if grupo and es_relevante:
        cliente = data.get('cliente', 'Sin cliente')
        pieza = "%s-%s-AR" % (prefijo, numero)

        if estado == 'RECHAZADO':
            msg = ("⚠️ *ENVÍO RECHAZADO*\n\n"
                   "👤 *Cliente:* %s\n"
                   "📦 *Pieza:* %s\n"
                   "📍 *Movimiento:* %s\n"
                   "🏢 *Planta:* %s\n"
                   "📅 *Fecha:* %s\n\n"
                   "_El destinatario rechazó la pieza. Revisar si requiere acción._"
                   % (cliente, pieza, ultimo['historia'], ultimo['planta'], ultimo['fecha']))
        elif estado == 'ENTREGADO':
            msg = ("✅ *Envío entregado*\n\n"
                   "👤 *Cliente:* %s\n"
                   "📦 *Pieza:* %s\n"
                   "🏢 *Planta:* %s\n"
                   "📅 *Fecha:* %s"
                   % (cliente, pieza, ultimo['planta'], ultimo['fecha']))
        else:
            msg = ("📦 *Novedad en envío*\n\n"
                   "👤 *Cliente:* %s\n"
                   "📦 *Pieza:* %s\n"
                   "📍 *Movimiento:* %s\n"
                   "🏢 *Planta:* %s\n"
                   "📅 *Fecha:* %s"
                   % (cliente, pieza, ultimo['historia'], ultimo['planta'], ultimo['fecha']))

        responder_whatsapp(grupo, msg)

    return 'con_novedad' if es_relevante else 'cambio_no_relevante'

@app.route('/revisar-correos', methods=['GET'])
def revisar_correos():
    contadores = {'error': 0, 'sin_resultados': 0, 'primera_vez': 0, 'sin_cambios': 0, 'con_novedad': 0}
    consultados = 0

    try:
        docs = list(db.collection('correos').stream())
    except Exception as e:
        print("Error leyendo correos:", str(e))
        return jsonify({"status": "error", "detalle": str(e)}), 500

    grupo = os.environ.get("GREEN_API_GROUP_ID")

    for doc in docs:
        data = doc.to_dict()
        # Solo las piezas con codigo cargado y todavia no finalizadas
        if not data.get('caPrefijo') or not data.get('caNumero') or data.get('caFinalizado'):
            continue
        consultados += 1
        resultado = procesar_pieza_correo(doc, data, grupo)
        contadores[resultado] = contadores.get(resultado, 0) + 1

    errores = contadores['error']
    # Si fallaron todas, probablemente cambio el sitio o nos estan bloqueando:
    # conviene enterarse en vez de que el chequeo quede roto en silencio.
    if grupo and consultados > 0 and errores == consultados:
        responder_whatsapp(
            grupo,
            "⚠️ *Seguimiento de correo*\nNo se pudo consultar ninguna pieza hoy (%d intento/s). "
            "Puede que Correo Argentino haya cambiado el sitio o esté bloqueando las consultas." % errores
        )

    return jsonify({
        "status": "ok",
        "consultados": consultados,
        "primera_consulta": contadores['primera_vez'],
        "con_novedad": contadores['con_novedad'],
        "cambio_no_relevante": contadores.get('cambio_no_relevante', 0),
        "sin_cambios": contadores['sin_cambios'],
        "errores": errores
    }), 200

# Se llama desde el navegador (index.html) justo despues de cargar un
# seguimiento nuevo, para no esperar hasta el proximo chequeo programado.
# Lleva CORS propio porque el resto del bot no lo necesita (solo lo pegan
# el webhook de WhatsApp y cron-job.org, ambos server-to-server).
@app.route('/revisar-correo-individual', methods=['GET'])
def revisar_correo_individual():
    doc_id = request.args.get('id')
    if not doc_id:
        resp = jsonify({"status": "error", "detalle": "Falta el parametro id"})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 400

    ref = db.collection('correos').document(doc_id)
    snap = ref.get()
    if not snap.exists:
        resp = jsonify({"status": "error", "detalle": "No existe ese seguimiento"})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 404

    data = snap.to_dict()
    if not data.get('caPrefijo') or not data.get('caNumero'):
        resp = jsonify({"status": "error", "detalle": "Ese registro no tiene codigo de Correo Argentino"})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 400

    grupo = os.environ.get("GREEN_API_GROUP_ID")
    resultado = procesar_pieza_correo(ref, data, grupo)

    resp = jsonify({"status": "ok", "resultado": resultado})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp, 200

# --- ESCRITOS CON IA ---
# Todo lo que es fijo (representacion legal, encabezado, formulas de cierre)
# se arma con texto literal, sin pasar por el modelo: asi los datos duros
# (nombres, CUIT, fechas) nunca corren riesgo de ser alterados por la IA.
# Claude solo redacta dos parrafos puntuales, donde hace falta traducir
# lenguaje informal a estilo formal: la descripcion de tareas y el relato de
# los hechos posteriores al despido.

BLOQUES_REPRESENTACION = {
    'ciardiello': (
        'acompañado por la Dra. Ciardiello María Paula M.P. 1-41227, CIDI NIVEL 2, '
        'CUIT 27-38003502-8, correo electrónico paulaciardiello@hotmail.com, '
        'cel. 3517541685, constituyendo domicilio a los efectos procesales en calle '
        'Arturo M. Bas Nº 389 1º Piso Oficina "A"'
    ),
    'casih': (
        'acompañado por el Dr. Pablo CASIH M.P. 1-31142 CIDI NIVEL 2, '
        'CUIT 20-22220869-7, correo electrónico pcasih@hotmail.com, '
        'cel. 3515102323, constituyendo domicilio a los efectos procesales en calle '
        'Arturo M. Bas Nº 389 1º Piso Oficina "A"'
    ),
    'ambos': (
        'acompañado por la Dra. Ciardiello María Paula M.P. 1-41227, CIDI NIVEL 2, '
        'CUIT 27-38003502-8, correo electrónico paulaciardiello@hotmail.com, '
        'cel. 3517541685 y por el Dr. Pablo CASIH M.P. 1-31142 CIDI NIVEL 2, '
        'CUIT 20-22220869-7, correo electrónico pcasih@hotmail.com, '
        'cel. 3515102323, constituyendo domicilio a los efectos procesales en calle '
        'Arturo M. Bas Nº 389 1º Piso Oficina "A"'
    ),
}

# El formulario solo soporta un denunciado (persona, empresa, o una persona
# titular de una empresa/comercio), asi que siempre es singular. Si hay
# persona fisica (sola o junto con una empresa), el sujeto es esa persona y
# se usa su genero. Si es solo empresa, decimos "la empresa" explicitamente
# en vez de "la denunciada" -- decir "la denunciada" ahi se lee como si
# fuera una persona, cuando en realidad es una razon social.
def denunciado_es_femenino(den):
    return (den.get('genero') or '').strip().lower() == 'femenino'

def de_denunciado(den):
    if not den.get('esPersona'):
        return 'de la empresa'
    return 'de la denunciada' if denunciado_es_femenino(den) else 'del denunciado'

def articulo_denunciado(den):
    if not den.get('esPersona'):
        return 'la empresa'
    return 'la denunciada' if denunciado_es_femenino(den) else 'el denunciado'

def armar_bloque_denunciado(den):
    persona = (den.get('esPersona') and (den.get('nombre') or '').strip())
    empresa = (den.get('esEmpresa') and (den.get('empresa') or '').strip())
    domicilio = (den.get('domicilio') or '').strip()
    ciudad = (den.get('ciudad') or '').strip()
    fantasia = (den.get('nombreFantasia') or '').strip()
    email = (den.get('email') or '').strip()

    # Persona fisica que opera un comercio con nombre de fantasia propio
    # (sin ser una empresa/CUIT aparte): usamos esta forma en vez de la de
    # "persona + empresa" para no pedirle un segundo CUIT que no existe.
    if persona and fantasia:
        base = 'mi empleador/a %s CUIL %s, titular y/o responsable del local comercial que gira bajo el nombre de fantasía: "%s"' % (
            den.get('nombre', ''), den.get('cuitPersona', ''), fantasia
        )
    elif persona and empresa:
        base = 'mi empleador/a %s CUIL %s, titular y/o responsable de la empresa %s CUIT %s, dedicada a %s' % (
            den.get('nombre', ''), den.get('cuitPersona', ''),
            den.get('empresa', ''), den.get('cuitEmpresa', ''), den.get('rubro', '')
        )
    elif persona:
        base = 'mi empleador/a %s CUIL %s' % (den.get('nombre', ''), den.get('cuitPersona', ''))
    elif empresa:
        base = 'la empresa %s CUIT %s, dedicada a %s' % (
            den.get('empresa', ''), den.get('cuitEmpresa', ''), den.get('rubro', '')
        )
    else:
        base = 'la parte denunciada'

    if domicilio or ciudad:
        base += ', con domicilio en %s%s' % (
            domicilio, (' de la ciudad de %s' % ciudad) if ciudad else ''
        )
    if email:
        base += ', correo electrónico %s' % email
    return base

def armar_bloque_registro(relacion, denunciado):
    texto = (
        "Que durante toda la relación laboral presté servicios de manera personal, "
        "habitual y bajo dependencia %s, sin que mi vínculo laboral "
        "se encontrara debidamente registrado ante los organismos correspondientes, "
        "en violación a la normativa laboral vigente."
    ) % de_denunciado(denunciado)
    extra = (relacion.get('registroParcial') or '').strip()
    if extra:
        texto += '\n\n' + extra
    return texto

# Relato fijo para el motivo "impedimento de ingreso" (el empleador no deja
# volver a trabajar). Otros motivos (accidente, falta de pago, etc.) van a
# necesitar su propio bloque el dia que se agreguen.
#
# La fecha del hecho queda en una oracion fija, separada, para que nunca
# pase por la IA (es un dato duro como el CUIT o un nombre). El resto del
# relato (que paso, como respondio el empleador) lo redacta la IA a partir
# de un texto libre que carga el usuario, asi que sirve para cualquier
# motivo de denuncia, no solo "impedimento de ingreso".
def armar_bloque_relato(relacion, parrafo_relato):
    fecha = (relacion.get('fechaDespido') or '').strip()
    partes = []
    if fecha:
        partes.append(
            'Que la relación laboral se desarrolló en los términos precedentemente expuestos '
            'hasta el día %s, fecha desde la cual se originan los hechos que motivan la presente '
            'denuncia.' % fecha
        )
    if parrafo_relato:
        partes.append(parrafo_relato)
    return '\n\n'.join(partes)

BLOQUE_CIERRE_RECLAMO = """Que conforme lo acontecido, las injurias ocasionadas, intimo a que me abone Liquidación Final, conforme la verdadera relación fáctica laboral y haberes adeudados y diferencias de haberes, SAC y Vacaciones No Gozadas por el plazo de prescripción, Indemnización por despido sin causa, antigüedad (art. 245 LCT), sustitutiva de preaviso (art. 232 LCT), Integración mes de despido (art. 233 LCT) y haga entrega de Certificaciones de Servicios y Remuneraciones estipuladas en Art. 80 LCT, sirviendo el presente como emplazamiento de dicha entrega."""

# --- REDACCION ASISTIDA POR IA ---
# Estos son los unicos tres puntos del escrito donde entra la IA. En todos
# los casos, los datos duros (fechas, numeros de CD, nombres) quedan afuera
# del texto que le pasamos a la IA: o bien se insertan aparte en la plantilla
# fija, o bien se le pide explicitamente que no invente ni repita datos.

SYSTEM_PROMPT_TAREAS = (
    "Sos un asistente de redacción para un estudio jurídico laboralista de Córdoba, Argentina. "
    "Vas a recibir una descripción informal de las tareas que hacía un trabajador y tenés que "
    "convertirla en UN párrafo formal, en primera persona, con el registro de una denuncia "
    "laboral ante el Ministerio de Trabajo, arrancando de forma natural (por ejemplo: 'Que las "
    "tareas desempeñadas por mi parte consistían en...').\n\n"
    "Reglas:\n"
    "- Nunca inventes datos que no estén en el texto que te pasaron.\n"
    "- Devolvé únicamente el párrafo final, sin comentarios antes o después."
)

SYSTEM_PROMPT_RELATO = (
    "Sos un asistente de redacción para un estudio jurídico laboralista de Córdoba, Argentina. "
    "Vas a recibir un relato informal de los hechos que motivan una denuncia laboral (impedimento "
    "de ingreso, despido directo, falta de pago, accidente de trabajo, u otro motivo) y tenés que "
    "convertirlo en uno o más párrafos formales, en primera persona, con el registro de una "
    "denuncia laboral ante el Ministerio de Trabajo, arrancando cada párrafo de forma natural "
    "(por ejemplo: 'Que...', 'Que, ante dicha situación...').\n\n"
    "Justo antes de tu texto, el escrito ya incluye una oración fija que dice que la relación "
    "laboral se desarrolló con normalidad hasta la fecha del hecho (la fecha ya quedó dicha ahí). "
    "Tu redacción tiene que continuar directamente contando qué pasó a partir de ese momento, sin "
    "repetir esa fecha ni reformular que 'la relación se desarrolló hasta tal fecha' — arrancá "
    "directo en el hecho en sí (ej: 'Que a partir de ese momento comencé a sufrir...').\n\n"
    "Reglas:\n"
    "- Nunca inventes datos que no estén en el texto que te pasaron.\n"
    "- No repitas fechas exactas salvo que te las hayan dado en el texto.\n"
    "- Devolvé únicamente el/los párrafo(s) final(es), sin comentarios antes o después."
)

SYSTEM_PROMPT_RESUMEN_RESPUESTA = (
    "Sos un asistente de redacción para un estudio jurídico laboralista de Córdoba, Argentina. "
    "Vas a recibir un resumen informal de lo que respondió la empresa/denunciado mediante una "
    "carta documento, y tenés que convertirlo en un fragmento formal, en discurso indirecto "
    "(nunca cites textual ni uses comillas), que continúe naturalmente la frase: 'Que, no obstante "
    "lo intimado, con fecha [FECHA] la parte denunciada remitió CD Nº [NUMERO], ...' — tu "
    "respuesta se pega justo después de esa coma.\n\n"
    "Ejemplo de respuesta esperada: 'mediante la cual reconoció la existencia de la relación "
    "laboral, aunque en disconformidad con la fecha de ingreso, e intimó erróneamente a mi parte "
    "a reintegrarse a sus tareas bajo apercibimiento de abandono, lo cual resulta a todas luces "
    "falaz, contradictorio y contrario a derecho.'\n\n"
    "Reglas:\n"
    "- Arrancá siempre con 'mediante la cual...' o una construcción equivalente que continúe la oración.\n"
    "- No repitas la fecha ni el número de CD: ya están en la frase anterior.\n"
    "- Nunca inventes datos que no estén en el texto que te pasaron.\n"
    "- Devolvé únicamente el fragmento final, sin comentarios antes o después."
)

SYSTEM_PROMPT_NORMALIZAR_MAYUSCULAS = (
    "Tu única tarea es corregir mayúsculas y minúsculas en un domicilio, sin cambiar nada más. "
    "Convertí el texto a la capitalización normal de una dirección en español: primera letra en "
    "mayúscula en cada palabra con sentido propio (calles, barrios, manzanas, lotes, nombres "
    "propios); artículos y preposiciones como 'de', 'del', 'la', 'los' en minúscula salvo que "
    "sean la primera palabra; abreviaturas de uso común (Nº, Bº, Mza., Dpto., Piso) tal como se "
    "escriben habitualmente.\n\n"
    "Reglas estrictas:\n"
    "- No agregues ni quites palabras, números, ni signos de puntuación.\n"
    "- No corrijas ortografía ni cambies ninguna palabra, solo su capitalización.\n"
    "- Devolvé únicamente el texto corregido, sin comentarios antes o después."
)

def redactar_con_ia(client, system_prompt, texto_informal):
    response = client.messages.create(
        model="claude-sonnet-5",
        max_tokens=1024,
        system=system_prompt,
        messages=[{"role": "user", "content": texto_informal}]
    )
    return "".join([b.text for b in response.content if b.type == "text"]).strip()

def normalizar_mayusculas_domicilio(client, texto):
    texto = (texto or '').strip()
    # Solo tocamos el texto si viene en mayusculas sostenidas (sin ninguna
    # minuscula) -- si ya tiene una capitalizacion razonable, lo dejamos tal
    # cual para no gastar una llamada a la IA de mas.
    if not texto or texto == texto.upper() and texto == texto.lower():
        return texto  # sin letras (solo numeros/simbolos), nada que normalizar
    if texto != texto.upper():
        return texto  # ya tiene minusculas, asumimos que esta bien escrito
    try:
        return redactar_con_ia(client, SYSTEM_PROMPT_NORMALIZAR_MAYUSCULAS, texto)
    except Exception as e:
        print("Error normalizando mayusculas de domicilio:", str(e))
        return texto

# Cada item de correspondencia se arma con una formula fija, en orden
# cronologico; el contenido transcripto (cita textual del telegrama) nunca
# pasa por la IA en este paso, solo se inserta tal cual lo cargo/leyo el
# usuario.
def armar_telegrama_nuestro(item, denunciado):
    fecha = (item.get('fechaEnvio') or '').strip()
    numero = (item.get('numero') or '').strip()
    contenido = (item.get('contenido') or '').strip()
    recepcion = (item.get('fechaRecepcion') or '').strip()

    if numero:
        intro = 'Que, con fecha %s, remití TCL CD Nº %s, en los siguientes términos:' % (fecha, numero)
    else:
        intro = 'Que, con fecha %s, remití TCL, en los siguientes términos:' % fecha

    texto = intro + '\n\n"%s"' % contenido
    if recepcion:
        texto += '\n\nDicha intimación fue debidamente recibida por %s con fecha %s conforme constancia del Correo Argentino.' % (articulo_denunciado(denunciado), recepcion)
    return texto

def armar_respuesta_empresa(resp):
    fecha = (resp.get('fechaEnvio') or '').strip()
    numero = (resp.get('numero') or '').strip()
    contenido = (resp.get('contenido') or '').strip()
    recepcion = (resp.get('fechaRecepcion') or '').strip()
    modo = (resp.get('modo') or 'citar').strip()

    texto = 'Que, no obstante lo intimado, con fecha %s la parte denunciada remitió CD' % fecha
    if numero:
        texto += ' Nº %s' % numero
    if recepcion:
        texto += ', recibida por mi parte el día %s' % recepcion

    if modo == 'resumir':
        # "contenido" ya viene redactado por la IA (ver redactar_con_ia con
        # SYSTEM_PROMPT_RESUMEN_RESPUESTA) como un fragmento en discurso
        # indirecto que continua esta misma oracion.
        texto += ', %s' % contenido
    else:
        texto += ', mediante la cual pretendió responder los reclamos, en los siguientes términos:\n\n"%s"' % contenido
    return texto

def armar_item_correspondencia(item, denunciado, alguna_respuesta_ref):
    if item.get('emisor') == 'empresa':
        alguna_respuesta_ref[0] = True
        return [armar_respuesta_empresa(item)]

    partes = [armar_telegrama_nuestro(item, denunciado)]
    if item.get('huboRespuesta') and item.get('respuesta'):
        partes.append(armar_respuesta_empresa(item['respuesta']))
        alguna_respuesta_ref[0] = True
    return partes

def armar_bloque_correspondencia(correspondencia, denunciado):
    partes = []
    alguna_respuesta_ref = [False]
    for item in correspondencia:
        partes.extend(armar_item_correspondencia(item, denunciado, alguna_respuesta_ref))
    if correspondencia and not alguna_respuesta_ref[0]:
        partes.append('Al día de la fecha %s no ha contestado ninguna de las intimaciones remitidas por esta parte.' % articulo_denunciado(denunciado))
    return '\n\n'.join(partes)

def armar_encabezado_cliente(cliente):
    nombre = (cliente.get('nombre') or '').strip()
    dni = (cliente.get('dni') or '').strip()
    cuit = (cliente.get('cuit') or '').strip()
    edad = (cliente.get('edad') or '').strip()
    estado_civil = (cliente.get('estadoCivil') or '').strip()
    nacionalidad = (cliente.get('nacionalidad') or '').strip()
    situacion = (cliente.get('situacionLaboral') or '').strip()
    domicilio = (cliente.get('domicilio') or '').strip()
    ciudad = (cliente.get('ciudad') or '').strip()
    telefono = (cliente.get('telefono') or '').strip()
    email = (cliente.get('email') or '').strip()
    bloque_rep = BLOQUES_REPRESENTACION.get(cliente.get('representacion'), BLOQUES_REPRESENTACION['ciardiello'])

    texto = nombre
    if dni:
        texto += ' DNI %s' % dni
    texto += ', CUIL %s' % cuit
    if edad:
        texto += ', %s años de edad' % edad
    texto += ', de estado civil %s, de nacionalidad %s' % (estado_civil, nacionalidad)
    if situacion:
        texto += ', actualmente %s' % situacion
    texto += (
        ', con domicilio real en %s, %s. Celular %s, correo electrónico: %s, %s, '
        'ambos de esta ciudad, ante S.S. respetuosamente comparezco y digo:'
    ) % (domicilio, ciudad, telefono, email, bloque_rep)
    return texto

# Si la fecha de ingreso es una fecha completa (dd/mm/aaaa), se le antepone
# "el" para que la oracion se lea bien ("...el 05/03/2025"). Si el usuario
# cargo otra cosa (ej. "en el mes de julio de 2020"), se deja tal cual la
# escribio -- ya trae su propio conector.
def formatear_fecha_ingreso(fecha):
    fecha = (fecha or '').strip()
    if re.match(r'^\d{1,2}/\d{1,2}/\d{2,4}$', fecha):
        return 'el %s' % fecha
    return fecha

def armar_escrito_denuncia_trabajo(cliente, denunciado, relacion, correspondencia, parrafo_tareas, parrafo_relato):
    encabezado_cliente = armar_encabezado_cliente(cliente)
    bloque_den = armar_bloque_denunciado(denunciado)
    bloque_registro = armar_bloque_registro(relacion, denunciado)
    bloque_relato = armar_bloque_relato(relacion, parrafo_relato)
    bloque_corr = armar_bloque_correspondencia(correspondencia, denunciado)

    texto = """FORMULA DENUNCIA.

Sr. Director de Reclamaciones Individuales
Departamento Provincial del Trabajo.
S___________________/____________D

%s

Que vengo a formular denuncia en contra de %s, comparecemos y decimos:

Que ingresé a trabajar bajo las órdenes %s %s, cumpliendo tareas acordes con la categoría %s.

Que mi jornada de trabajo era de %s de %s.

%s

%s

%s

%s

%s

Por los motivos expuestos, se solicita la intervención de esta Autoridad de Aplicación.

En definitiva, ante el incumplimiento de la normativa vigente por parte del Empleador, solicitó que el mismo sea citado, quien deberá concurrir con toda la documentación laboral, legajo personal, además de las constancias de pago de los aportes y contribuciones sindicales, sociales y previsionales, y los importes correspondientes a Liquidación Final e indemnizaciones de ley, a cuyo fin deberá fijar día y hora de audiencia de conciliación. Sin otro particular. Saludo a Ud. muy atte.""" % (
        encabezado_cliente,
        bloque_den,
        de_denunciado(denunciado), formatear_fecha_ingreso(relacion.get('fechaIngreso', '')), relacion.get('categoria', ''),
        relacion.get('jornada', ''), relacion.get('horarios', ''),
        parrafo_tareas,
        bloque_registro,
        bloque_relato,
        bloque_corr,
        BLOQUE_CIERRE_RECLAMO
    )
    # Si algun bloque quedo vacio (ej. sin correspondencia cargada), evita que
    # queden lineas en blanco de mas entre parrafos.
    return re.sub(r'\n{3,}', '\n\n', texto)

# --- EXPORTAR A WORD (.docx) ---
# Toma el texto ya generado (y eventualmente editado a mano en el textarea)
# y lo pasa a un documento .docx con tipografia, tamano, justificado y
# sangria acordes al modelo aprobado por el estudio, listo para retocar en
# Word antes de exportar a PDF.

FIRMAS_ABOGADOS = {
    'ciardiello': [('Ma. Paula Ciardiello', 'Abogada', 'M.P. 1-41227')],
    'casih': [('Pablo Casih', 'Abogado', 'M.P. 1-31142')],
    'ambos': [('Ma. Paula Ciardiello', 'Abogada', 'M.P. 1-41227'), ('Pablo Casih', 'Abogado', 'M.P. 1-31142')],
}

def _agregar_firma(celda, nombre, subt1='', subt2=''):
    celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    celda.paragraphs[0].add_run('_______________________')
    p_nombre = celda.add_paragraph(nombre)
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_nombre.runs:
        p_nombre.runs[0].bold = True
    for sub in (subt1, subt2):
        if sub:
            p_sub = celda.add_paragraph(sub)
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

def construir_docx_denuncia(texto, cliente):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    estilo = doc.styles['Normal']
    estilo.font.name = 'Times New Roman'
    estilo.font.size = Pt(12)
    estilo.paragraph_format.line_spacing = 1.15
    estilo.paragraph_format.space_after = Pt(0)

    marcador_cierre = 'Sin otro particular. Saludo a Ud. muy atte.'
    bloques = [b.strip('\n') for b in texto.split('\n\n') if b.strip()]

    for bloque in bloques:
        if bloque.strip() == 'FORMULA DENUNCIA.':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('FORMULA DENUNCIA.').bold = True
            continue

        if bloque.startswith('Sr. Director'):
            p = doc.add_paragraph()
            for i, linea in enumerate(bloque.split('\n')):
                if i > 0:
                    p.add_run().add_break()
                p.add_run(linea)
            continue

        if bloque.startswith('"') and bloque.endswith('"'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.right_indent = Cm(1.5)
            p.add_run(bloque).italic = True
            continue

        if marcador_cierre in bloque:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            antes, _, despues = bloque.partition(marcador_cierre)
            if antes:
                p.add_run(antes)
            p.add_run(marcador_cierre + despues).bold = True
            continue

        p = doc.add_paragraph(bloque)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_paragraph()
    doc.add_paragraph()

    abogados = FIRMAS_ABOGADOS.get(cliente.get('representacion'), FIRMAS_ABOGADOS['ciardiello'])
    tabla = doc.add_table(rows=1, cols=1 + len(abogados))
    celdas = tabla.rows[0].cells

    nombre_cliente = (cliente.get('nombre') or '').strip()
    dni_cliente = (cliente.get('dni') or '').strip()
    _agregar_firma(celdas[0], nombre_cliente, ('DNI %s' % dni_cliente) if dni_cliente else '')
    for i, (nombre_ab, titulo_ab, mp_ab) in enumerate(abogados):
        _agregar_firma(celdas[1 + i], nombre_ab, titulo_ab, mp_ab)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/generar-escrito-docx', methods=['POST', 'OPTIONS'])
def generar_escrito_docx():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    body = request.get_json(silent=True) or {}
    texto = (body.get('texto') or '').strip()
    cliente = body.get('cliente') or {}

    if not texto:
        resp = jsonify({"status": "error", "detalle": "Falta el texto del escrito."})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 400

    try:
        buffer = construir_docx_denuncia(texto, cliente)
    except Exception as e:
        print("Error generando docx:", str(e))
        resp = jsonify({"status": "error", "detalle": "No se pudo generar el archivo Word."})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 500

    nombre_archivo = 'Denuncia - %s.docx' % ((cliente.get('nombre') or 'escrito').strip() or 'escrito')

    resp = send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=nombre_archivo
    )
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

# --- TELEGRAMA (TCL) SOBRE EL FORMULARIO DE CORREO ARGENTINO ---
# El modelo de Word del estudio escribia los datos como parrafos comunes
# separados con espacios para simular las dos columnas. Como el formulario
# esta anclado y el texto fluye, los datos terminaban encima de las lineas
# impresas y las columnas se desbordaban al cambiar el largo de un dato.
# Aca cada dato va en su propia caja anclada a la pagina, en coordenadas
# absolutas medidas sobre el formulario, asi siempre cae donde corresponde.

TG_CM_EMU = 360000
TG_CM_PT = 28.3465

# Geometria del formulario, en cm respecto del borde de la hoja.
#
# La plantilla esta normalizada: todo el formulario quedo anclado a la pagina.
# En el modelo original venia partido en tres sistemas de coordenadas — una
# parte anclada a la pagina y otras dos ancladas a parrafos — y como cada
# programa calcula distinto el alto del texto anterior, en Word, Pages y
# Google Docs el encabezado se corria y el recuadro del cuerpo se subia dos
# centimetros tapando la fila de Localidad/Provincia.
TG_LINEA = {'fila1': 4.14, 'fila2': 5.15, 'fila3': 6.16, 'fila4': 7.18}
TG_IZQ_X, TG_IZQ_AN = 1.09, 8.89
TG_DER_X, TG_DER_AN = 10.98, 8.91
TG_CUERPO = dict(x=1.09, y=8.60, an=18.87, al=16.57)
TG_ALTO_CAMPO = 0.50
# Banda libre entre la fila 4 y el recuadro (debajo de las etiquetas
# Localidad/Provincia): ahi entran el CUIL y la leyenda de mas de 30 palabras.
TG_BANDA_LIBRE = 8.05

PLANTILLA_TELEGRAMA = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), 'plantillas', 'modelo_telegrama.docx')


def _tg_emu(cm):
    return int(round(cm * TG_CM_EMU))


def _tg_caja(idx, x_cm, y_cm, an_cm, al_cm, texto, tam_pt=10, negrita=False,
             alineacion='left', anclaje='b', relh='page', relv='page'):
    """Caja de texto anclada, sin borde ni relleno.

    anclaje 'b' apoya el texto en el borde inferior (para que quede justo
    sobre la linea del formulario); 't' lo arranca desde arriba.

    relh/relv definen respecto de que se posiciona. El formulario esta partido
    en dos sistemas: la mitad de abajo esta anclada a la pagina y la de arriba
    al parrafo, asi que cada dato tiene que usar el mismo que su parte del
    formulario o se despegan al abrirlo en otro programa.
    """
    sz = int(round(tam_pt * 2))  # w:sz va en medios puntos
    parrafos = []
    for linea in (texto or '').split('\n'):
        parrafos.append(
            '<w:p><w:pPr>'
            '<w:spacing w:after="0" w:before="0" w:line="240" w:lineRule="auto"/>'
            '<w:ind w:left="0" w:right="0" w:firstLine="0"/>'
            '<w:jc w:val="%s"/></w:pPr>'
            '<w:r><w:rPr>'
            '<w:rFonts w:ascii="Arial" w:cs="Arial" w:eastAsia="Arial" w:hAnsi="Arial"/>'
            '<w:sz w:val="%d"/><w:szCs w:val="%d"/>%s<w:rtl w:val="0"/></w:rPr>'
            '<w:t xml:space="preserve">%s</w:t></w:r></w:p>'
            % (alineacion, sz, sz,
               '<w:b w:val="1"/><w:bCs w:val="1"/>' if negrita else '',
               xml_escape(linea))
        )
    cx, cy = _tg_emu(an_cm), _tg_emu(al_cm)
    return (
        '<w:r><w:drawing><wp:anchor allowOverlap="1" behindDoc="0" distB="0" distT="0"'
        ' distL="0" distR="0" hidden="0" layoutInCell="1" locked="0" relativeHeight="%d"'
        ' simplePos="0">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="%s"><wp:posOffset>%d</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="%s"><wp:posOffset>%d</wp:posOffset></wp:positionV>'
        '<wp:extent cx="%d" cy="%d"/>'
        '<wp:effectExtent b="0" l="0" r="0" t="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="%d" name="campo%d"/>'
        '<a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvSpPr txBox="1"/>'
        '<wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="%d" cy="%d"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom><a:noFill/>'
        '<a:ln><a:noFill/></a:ln></wps:spPr>'
        '<wps:txbx><w:txbxContent>%s</w:txbxContent></wps:txbx>'
        '<wps:bodyPr anchor="%s" anchorCtr="0" bIns="0" lIns="0" rIns="0" tIns="0"'
        ' spcFirstLastPara="0" wrap="square"><a:noAutofit/></wps:bodyPr>'
        '</wps:wsp></a:graphicData></a:graphic>'
        '</wp:anchor></w:drawing></w:r>'
        % (900 + idx, relh, _tg_emu(x_cm), relv, _tg_emu(y_cm), cx, cy,
           9000 + idx, idx, cx, cy, ''.join(parrafos), anclaje)
    )


def _tg_ajustar_tam(texto, an_cm, tam_pt, minimo=6.0):
    """Achica la tipografia lo justo para que el dato entre en su renglon."""
    texto = (texto or '').strip()
    if not texto:
        return tam_pt
    # Ancho medio de caracter en Arial, en "em": las mayusculas son bastante
    # mas anchas, y estos campos suelen cargarse en mayusculas.
    letras = [c for c in texto if c.isalpha()]
    mayus = (sum(1 for c in letras if c.isupper()) / len(letras)) if letras else 0
    em = 0.73 if mayus > 0.6 else 0.58
    necesario = (an_cm * TG_CM_PT * 0.97) / (len(texto) * em)
    return max(minimo, min(tam_pt, round(necesario, 1)))


def _tg_sobre_linea(idx, x_cm, y_linea, an_cm, texto, tam_pt=10, **kw):
    """Caja que apoya el texto justo por encima de una linea del formulario."""
    return _tg_caja(idx, x_cm, y_linea - TG_ALTO_CAMPO, an_cm, TG_ALTO_CAMPO,
                    texto, tam_pt=_tg_ajustar_tam(texto, an_cm, tam_pt), **kw)


def construir_cajas_telegrama(d):
    """Devuelve el XML de todas las cajas, en coordenadas de pagina."""
    c = []
    n = iter(range(1, 999))

    # Fila 1: razon social del destinatario / apellido y nombre del remitente
    c.append(_tg_sobre_linea(next(n), TG_IZQ_X, TG_LINEA['fila1'], TG_IZQ_AN,
                             d.get('destNombre', ''), negrita=True))
    c.append(_tg_sobre_linea(next(n), TG_DER_X, TG_LINEA['fila1'], TG_DER_AN,
                             d.get('remNombre', ''), negrita=True))

    # Fila 2: ramo o actividad principal / DNI del remitente
    c.append(_tg_sobre_linea(next(n), TG_IZQ_X, TG_LINEA['fila2'], TG_IZQ_AN,
                             d.get('destRamo', ''), tam_pt=9, negrita=True))
    c.append(_tg_sobre_linea(next(n), TG_DER_X, TG_LINEA['fila2'], TG_DER_AN,
                             d.get('remDni', ''), negrita=True))

    # "Más de 30 palabras" es una anotacion del estudio que en el modelo pisa
    # el renglon del ramo; la reubicamos en el hueco libre entre el banner
    # azul "Telegrama Ley Nº 23.789" y el banner verde "DESTINATARIO",
    # alineada a la izquierda con ambos.
    c.append(_tg_caja(next(n), TG_IZQ_X, 2.05, TG_IZQ_AN, 0.55,
                      'Más de 30 palabras', tam_pt=10, negrita=True))

    # Fila 3: domicilio + codigo postal (el CP va al final de cada columna)
    an_cp = 2.40
    an_dom = 6.30
    c.append(_tg_sobre_linea(next(n), TG_IZQ_X, TG_LINEA['fila3'], an_dom,
                             d.get('destDomicilio', ''), tam_pt=9, negrita=True))
    c.append(_tg_sobre_linea(next(n), TG_IZQ_X + TG_IZQ_AN - an_cp, TG_LINEA['fila3'],
                             an_cp, d.get('destCp', ''), tam_pt=9, alineacion='center', negrita=True))
    c.append(_tg_sobre_linea(next(n), TG_DER_X, TG_LINEA['fila3'], an_dom,
                             d.get('remDomicilio', ''), tam_pt=9, negrita=True))
    c.append(_tg_sobre_linea(next(n), TG_DER_X + TG_DER_AN - an_cp, TG_LINEA['fila3'],
                             an_cp, d.get('remCp', ''), tam_pt=9, alineacion='center', negrita=True))

    # Fila 4: localidad + provincia (mitad y mitad de cada columna)
    mitad = TG_IZQ_AN / 2 - 0.15
    c.append(_tg_sobre_linea(next(n), TG_IZQ_X, TG_LINEA['fila4'], mitad,
                             d.get('destLocalidad', ''), tam_pt=9, negrita=True))
    c.append(_tg_sobre_linea(next(n), TG_IZQ_X + TG_IZQ_AN / 2, TG_LINEA['fila4'], mitad,
                             d.get('destProvincia', ''), tam_pt=9, alineacion='center', negrita=True))
    c.append(_tg_sobre_linea(next(n), TG_DER_X, TG_LINEA['fila4'], mitad,
                             d.get('remLocalidad', ''), tam_pt=9, negrita=True))
    c.append(_tg_sobre_linea(next(n), TG_DER_X + TG_DER_AN / 2, TG_LINEA['fila4'], mitad,
                             d.get('remProvincia', ''), tam_pt=9, alineacion='center', negrita=True))

    # CUIL del destinatario: espacio libre entre la fila 4 y el recuadro
    cuil = (d.get('destCuil') or '').strip()
    if cuil:
        if not cuil.upper().startswith('CUIL'):
            cuil = 'CUIL ' + cuil
        c.append(_tg_caja(next(n), TG_IZQ_X, TG_BANDA_LIBRE, 9.0, 0.55,
                          cuil, tam_pt=9, negrita=True))

    # Cuerpo del telegrama, dentro del recuadro
    cuerpo = d.get('cuerpo', '') or ''
    # si el texto es muy largo, achicamos un punto para que entre en el recuadro
    tam_cuerpo = 9 if len(cuerpo) <= 3600 else 8
    c.append(_tg_caja(next(n), TG_CUERPO['x'] + 0.20, TG_CUERPO['y'] + 0.18,
                      TG_CUERPO['an'] - 0.40, TG_CUERPO['al'] - 0.30,
                      cuerpo, tam_pt=tam_cuerpo, alineacion='both', anclaje='t'))
    return ''.join(c)


_TG_RE_P = re.compile(r'<w:p(?=[ >/])|</w:p>')


def _tg_fin_parrafo(doc, indice):
    """Posicion del </w:p> que cierra el parrafo de nivel superior nro `indice`.

    No sirve contar "</w:p>" a secas: el formulario tiene cajas de texto con
    parrafos adentro, asi que hay que llevar la cuenta de anidamiento y contar
    solo los que cierran a nivel del body.
    """
    inicio = doc.index('<w:body>')
    nivel = 0
    n = 0
    for m in _TG_RE_P.finditer(doc, inicio):
        if m.group(0) == '</w:p>':
            nivel -= 1
            if nivel == 0:
                if n == indice:
                    return m.start()
                n += 1
        elif doc[m.start():m.end() + 40].split('>')[0].endswith('/'):
            if nivel == 0:            # <w:p/> vacio de nivel superior
                if n == indice:
                    return m.start()
                n += 1
        else:
            nivel += 1
    raise ValueError('No se encontro el parrafo %d del modelo.' % indice)


def generar_telegrama_docx(datos):
    """Devuelve los bytes de un .docx: formulario oficial + datos ubicados."""
    with open(PLANTILLA_TELEGRAMA, 'rb') as f:
        modelo = f.read()
    zin = zipfile.ZipFile(BytesIO(modelo))
    doc = zin.read('word/document.xml').decode('utf-8')

    # "Más de 30 palabras" viene como texto que fluye y cae justo sobre el
    # renglon del ramo; lo sacamos del flujo y se repone como caja ubicada.
    doc = doc.replace('<w:t xml:space="preserve">Más de 30 palabras</w:t>',
                      '<w:t xml:space="preserve"></w:t>')

    corte = _tg_fin_parrafo(doc, 0)
    doc = doc[:corte] + construir_cajas_telegrama(datos) + doc[corte:]

    out = BytesIO()
    with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                data = doc.encode('utf-8')
            zout.writestr(item, data)
    out.seek(0)
    return out.read()


# El cuerpo se arma con texto fijo + los datos cargados. La IA solo redacta
# los tramos que vienen en lenguaje informal (jornada y observaciones), igual
# que en la denuncia: los datos duros nunca pasan por el modelo.
CIERRE_TELEGRAMA = (
    'La intimación efectuada es bajo apercibimiento del reclamo de daños y perjuicios '
    'por el perjuicio irreparable al estar la relación laboral de manera informal. Hago '
    'reserva de ampliar argumentos en la etapa correspondiente. Todo bajo apercibimiento '
    'de considerar su silencio o desconocimiento insincero de las circunstancias fácticas '
    'de la relación laboral como grave injuria a mis derechos laborales y colocarme en '
    'situación de despido indirecto. Hago reservas de salarios caídos. A los fines del '
    'presente constituyo domicilio legal en Arturo M. Bas Nº 389 Piso 1º Of. “A”, Ciudad '
    'de Córdoba “ESTUDIO JURIDICO CASIH Y ASOC.” Tel. 3516327201. Queda Uds. debidamente '
    'notificados y constituidos en mora.'
)


def armar_cuerpo_telegrama(d, jornada_redactada, observaciones_redactadas):
    fem = (d.get('remGenero') or '').strip().lower() == 'femenino'
    despedido = 'despedida' if fem else 'despedido'

    encabezado = 'EN CARÁCTER DE TITULAR'
    descripcion = (d.get('destDescripcion') or '').strip()
    if descripcion:
        encabezado += ' DE %s' % descripcion
    fantasia = (d.get('destFantasia') or '').strip()
    if fantasia:
        encabezado += ' QUE GIRA BAJO EL NOMBRE DE FANTASIA : “%s”' % fantasia
    extension = (d.get('extension') or '').strip()
    if extension:
        encabezado += ' - %s' % extension

    partes = []
    partes.append(
        '%s: Atento al impedimento injustificado que existe a que ingrese a mi lugar '
        'habitual de trabajo, desde el día %s, dando respuestas dilatorias y evasivas, '
        'negándose a asignarme las tareas propias de mi categoría y especialidad, los '
        'EMPLAZO para que en el término de dos (2) días hábiles aclaren fehacientemente '
        'mi situación laboral permitiéndome el ingreso a mi lugar de trabajo y '
        'proporcionándome las tareas propias de mi categoría en lugar y horario habitual '
        'de trabajo, bajo apercibimiento de considerar su silencio o desconocimiento una '
        'injuria grave y considerarme %s por su exclusiva culpa.'
        % (encabezado, d.get('fechaImpedimento', ''), despedido)
    )
    partes.append(
        'Asimismo, a no tener constancia que mi relación de dependencia, económica, '
        'técnica y jurídica para con Ud. se encuentre debidamente registrada ante los '
        'organismos pertinentes, a pesar de mis pedidos en tal sentido, lo EMPLAZO E '
        'INTIMO para que en el plazo de treinta (30) días proceda a registrarla en forma '
        'correcta, a cuyo fin aporto los siguientes datos:'
    )

    datos = []
    def agregar(etiqueta, valor, sep='.'):
        valor = (valor or '').strip()
        if valor:
            datos.append('%s: %s%s' % (etiqueta, valor, sep))
    agregar('Nombre y Apellido', d.get('remNombre', ''), '')
    agregar('DNI', d.get('remDni', ''))
    edad = (d.get('remEdad') or '').strip()
    if edad:
        datos.append('Edad: %s años.' % edad)
    agregar('Fecha de Nacimiento', d.get('remNacimiento', ''), ';')
    agregar('Estado Civil', d.get('remEstadoCivil', ''))
    agregar('Nacionalidad', d.get('remNacionalidad', ''))
    agregar('Domicilio', d.get('remDomicilioReal', ''))
    agregar('Fecha de ingreso', d.get('fechaIngreso', ''))
    agregar('Categoría', d.get('categoria', ''))
    if jornada_redactada:
        datos.append('Jornada de Trabajo: %s' % jornada_redactada.rstrip('.') + '.')
    partes.append(' '.join(datos))

    if observaciones_redactadas:
        partes.append(observaciones_redactadas)
    partes.append(CIERRE_TELEGRAMA)
    return ' '.join(p.strip() for p in partes if p and p.strip())


SYSTEM_PROMPT_TG_JORNADA = (
    "Sos un asistente de redacción para un estudio jurídico laboralista de Córdoba, "
    "Argentina. Vas a recibir la descripción informal de la jornada de trabajo de una "
    "persona y tenés que reescribirla en el registro formal de un telegrama laboral "
    "(Ley 23.789), como una frase corrida que continúa la oración 'Jornada de Trabajo: '.\n\n"
    "Ejemplo de salida: 'días fijos Martes, Viernes y Sábado, cubriendo además francos "
    "cuando era requerido, con horarios rotativos, cumpliendo turnos de mañana de 09:30 a "
    "13:30 horas y turnos de tarde de 17:00 a 21:00 horas'\n\n"
    "Reglas:\n"
    "- No agregues 'Jornada de Trabajo:' al principio, eso ya está.\n"
    "- Nunca inventes días, horarios ni datos que no estén en el texto.\n"
    "- Devolvé únicamente la frase, sin comentarios antes o después."
)

SYSTEM_PROMPT_TG_OBSERVACIONES = (
    "Sos un asistente de redacción para un estudio jurídico laboralista de Córdoba, "
    "Argentina. Vas a recibir una observación informal sobre la relación laboral y tenés "
    "que convertirla en una o dos oraciones formales, en primera persona, para incluir en "
    "un telegrama laboral (Ley 23.789), arrancando de forma natural (por ejemplo: "
    "'Asimismo, hago saber que...').\n\n"
    "Reglas:\n"
    "- Nunca inventes datos que no estén en el texto que te pasaron.\n"
    "- Devolvé únicamente el texto final, sin comentarios antes o después."
)


@app.route('/generar-telegrama', methods=['POST', 'OPTIONS'])
def generar_telegrama():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    d = request.get_json(silent=True) or {}

    def error(msg, code=400):
        resp = jsonify({"status": "error", "detalle": msg})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, code

    if not (d.get('destNombre') or '').strip():
        return error('Falta el nombre o razón social del destinatario.')
    if not (d.get('remNombre') or '').strip():
        return error('Falta el nombre del remitente (el cliente).')
    if not (d.get('fechaImpedimento') or '').strip():
        return error('Falta la fecha desde la que se le impide el ingreso.')

    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return error('El servidor no tiene configurada la clave de la API de Claude.', 500)

    try:
        client = anthropic.Anthropic(api_key=api_key)
        jornada = ''
        if (d.get('jornada') or '').strip():
            jornada = redactar_con_ia(client, SYSTEM_PROMPT_TG_JORNADA, d['jornada'])
        observaciones = ''
        if (d.get('observaciones') or '').strip():
            observaciones = redactar_con_ia(
                client, SYSTEM_PROMPT_TG_OBSERVACIONES, d['observaciones'])
    except Exception as e:
        print("Error redactando telegrama con IA:", str(e))
        return error('No se pudo generar el telegrama. Intentá de nuevo en un momento.', 500)

    texto = armar_cuerpo_telegrama(d, jornada, observaciones)
    resp = jsonify({"status": "ok", "texto": texto})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp, 200


@app.route('/generar-telegrama-docx', methods=['POST', 'OPTIONS'])
def generar_telegrama_docx_endpoint():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    d = request.get_json(silent=True) or {}
    if not (d.get('cuerpo') or '').strip():
        resp = jsonify({"status": "error", "detalle": "Falta el texto del telegrama."})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 400

    try:
        contenido = generar_telegrama_docx(d)
    except Exception as e:
        print("Error generando telegrama docx:", str(e))
        resp = jsonify({"status": "error", "detalle": "No se pudo generar el archivo Word."})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 500

    nombre = (d.get('remNombre') or 'telegrama').strip() or 'telegrama'
    resp = send_file(
        BytesIO(contenido),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='Telegrama - %s.docx' % nombre
    )
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp


# ─── Telegrama de despido indirecto ───
# Mismo formulario físico que el telegrama de "aclare situación" (se reutiliza
# generar_telegrama_docx / construir_cajas_telegrama sin cambios), pero con un
# cuerpo distinto: acusa recibo de un TCL previo no contestado y notifica el
# despido indirecto ya efectivizado, en lugar de solo intimar.

TG_DESPIDO_INTIMACION_LCT = (
    'INTIMOLE para que en el plazo perentorio de cuatro (4) días hábiles abonen las '
    'indemnizaciones de ley: Liquidación final, Indemnización por despido indirecto, '
    'antigüedad (art. 245 LCT), sustitutiva de preaviso (art. 232 LCT), Integración de '
    'mes de despido (art. 233 LCT) con incidencia de horas extraordinarias y SAC, horas '
    'extraordinarias, diferencias de haberes, SAC y Vacaciones no gozadas por el tiempo '
    'no prescripto.'
)

def _tg_despido_intimacion_construccion(categoria, jornada):
    return (
        'Intimo tres días hábiles pago Liquidación Final debiendo incluir en ella los '
        'jornales caídos, adicional por presentismo jamás abonado del que soy acreedor '
        'por haber mantenido todo el vínculo asistencia perfecta, las diferencias de '
        'haberes a mi favor por abonar Uds. menos horas de las efectivamente trabajadas '
        'teniendo en cuenta mi real jornada de trabajo de %s, cumpliendo tareas de %s '
        'según CCT 76/75, todo ello bajo apercibimiento Art. 19 Ley 22.250. Intimo en dos '
        'días hábiles restitución Libreta de Aportes Ley 22.250 y pongan a disposición '
        'las sumas en concepto de Fondo Cese Laboral, bajo apercibimiento Art. 18 Ley '
        '22.250. Hago reservas de responsabilizar solidariamente a su comitente.'
        % (jornada, categoria)
    )

TG_DESPIDO_NO_REGISTRADO = (
    'Todo ello teniendo especialmente en consideración que la relación laboral se '
    'desarrolló íntegramente sin registración alguna ante los organismos '
    'correspondientes, en incumplimiento de las obligaciones legales a cargo del '
    'empleador, se reclamará el pago de daños y perjuicios por el perjuicio irreparable '
    'derivado de la informalidad de la relación laboral y su falta de registro ante los '
    'organismos correspondientes.'
)

TG_DESPIDO_DIF_APORTES = (
    'Por otro lado al tener constancia que las sumas retenidas de mis haberes en '
    'concepto de Jubilación, Ley 19032, Obra Social, ANSSAL, Cuota Sindical y Seguro de '
    'Vida no han sido efectivamente ingresadas a los organismos pertinentes lo intimo '
    'para que, dentro del término de TREINTA (30) días corridos contados a partir de la '
    'recepción de la presente, ingrese los importes adeudados, más los intereses y '
    'multas que pudieren corresponder, a los respectivos Organismos recaudadores, bajo '
    'apercibimiento del Art. 132 bis de la LCT.'
)

CIERRE_TELEGRAMA_DESPIDO = (
    'Ratifico domicilio legal en calle Arturo M. Bas Nº 389 Piso 1º Of. “A”, Ciudad de '
    'Córdoba “ESTUDIO JURIDICO CASIH Y ASOC.” Te. 3516327201. Queda Ud. debidamente '
    'notificad%s y constituid%s en mora.'
)


def armar_apertura_despido_indirecto(d, relato_redactado):
    descripcion = (d.get('destDescripcion') or '').strip()
    fantasia = (d.get('destFantasia') or '').strip()
    encabezado = ''
    if descripcion or fantasia:
        encabezado = 'EN CARÁCTER DE TITULAR'
        if descripcion:
            encabezado += ' DE %s' % descripcion
        if fantasia:
            encabezado += ' QUE GIRA BAJO EL NOMBRE DE FANTASIA: “%s”' % fantasia
        extension = (d.get('extension') or '').strip()
        if extension:
            encabezado += ' - %s' % extension
        encabezado += ': '

    numero = (d.get('tclNumero') or '').strip()
    fecha_envio = (d.get('tclFechaEnvio') or '').strip()
    fecha_recepcion = (d.get('tclFechaRecepcion') or '').strip()
    ref_tcl = 'Encontrándose vencidos los términos de mi TLC CD Nº %s' % numero
    if fecha_envio:
        ref_tcl += ' de fecha %s' % fecha_envio
    if fecha_recepcion:
        ref_tcl += ' entregado con fecha %s' % fecha_recepcion
    ref_tcl += ', sin aclarar mi situación laboral'

    fem_rem = (d.get('remGenero') or '').strip().lower() == 'femenino'
    despedido = 'despedida' if fem_rem else 'despedido'

    cierre_relato = (
        ', le comunico que todo ello me ocasiona una injuria laboral tan grave que '
        'imposibilita la prosecución del vínculo laboral, y conforme lo emplazado, hago '
        'efectivo el apercibimiento cursado por lo cual me considero %s por su exclusiva '
        'culpa.' % despedido
    )

    relato = (relato_redactado or '').strip().rstrip('.')
    cuerpo = ref_tcl
    if relato:
        cuerpo += ', %s' % relato
    cuerpo += cierre_relato

    texto = encabezado + cuerpo
    if not encabezado:
        texto = texto[0].upper() + texto[1:]
    return texto


def armar_cuerpo_despido_indirecto(d, relato_redactado):
    fem_dest = (d.get('destGenero') or '').strip().lower() == 'femenino'
    sufijo = 'a' if fem_dest else 'o'

    partes = [armar_apertura_despido_indirecto(d, relato_redactado)]

    regimen = (d.get('regimen') or 'lct').strip().lower()
    if regimen == 'construccion':
        partes.append(_tg_despido_intimacion_construccion(
            d.get('categoria', ''), d.get('jornada', '')))
    else:
        partes.append(TG_DESPIDO_INTIMACION_LCT)

    if d.get('noRegistrado'):
        partes.append(TG_DESPIDO_NO_REGISTRADO)
    if d.get('difAportes'):
        partes.append(TG_DESPIDO_DIF_APORTES)

    partes.append(CIERRE_TELEGRAMA_DESPIDO % (sufijo, sufijo))

    return ' '.join(p.strip() for p in partes if p and p.strip())


SYSTEM_PROMPT_TG_INJURIA_DESPIDO = (
    "Sos un asistente de redacción para un estudio jurídico laboralista de Córdoba, "
    "Argentina. Vas a recibir la descripción informal de los motivos por los que un "
    "cliente se considera despedido en forma indirecta (impedimento de ingreso, falta de "
    "tareas, respuestas evasivas, falta de registración, falta de pago de haberes, etc.) "
    "y tenés que convertirla en una frase corrida, en primera persona, que continúa SIN "
    "CORTE la oración: 'sin aclarar mi situación laboral, [TU TEXTO], le comunico que "
    "todo ello me ocasiona...'.\n\n"
    "Ejemplo de salida (para 'no me dejan entrar a trabajar hace un mes y no contestan "
    "nada'): 'siendo que continúa impidiéndome el ingreso a mi lugar de trabajo, se sigue "
    "negando a otorgarme tareas, con respuestas evasivas'\n\n"
    "Reglas:\n"
    "- Arrancá en minúscula y sin punto final: tu frase se inserta en medio de una "
    "oración más larga que ya arranca en mayúscula y sigue después con otro tramo fijo.\n"
    "- No repitas 'sin aclarar mi situación laboral', eso ya está antes de tu texto.\n"
    "- No agregues 'le comunico que...' al final, eso ya está después.\n"
    "- Nunca inventes hechos, fechas ni datos que no estén en el texto que te pasaron.\n"
    "- Devolvé únicamente la frase, sin comentarios antes o después."
)

SYSTEM_PROMPT_TG_JORNADA_DESPIDO = (
    "Sos un asistente de redacción para un estudio jurídico laboralista de Córdoba, "
    "Argentina. Vas a recibir la descripción informal de la jornada de trabajo de una "
    "persona del régimen de la construcción y tenés que reescribirla en el registro "
    "formal de un telegrama laboral, como una frase corrida que continúa la oración "
    "'teniendo en cuenta mi real jornada de trabajo de '.\n\n"
    "Ejemplo de salida: 'Lunes a Viernes de 08:00 a 17:00 hs'\n\n"
    "Reglas:\n"
    "- No agregues 'jornada de trabajo de' al principio, eso ya está.\n"
    "- Nunca inventes días, horarios ni datos que no estén en el texto.\n"
    "- Devolvé únicamente la frase, sin comentarios antes o después."
)


@app.route('/generar-telegrama-despido-indirecto', methods=['POST', 'OPTIONS'])
def generar_telegrama_despido_indirecto():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    d = request.get_json(silent=True) or {}

    def error(msg, code=400):
        resp = jsonify({"status": "error", "detalle": msg})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, code

    if not (d.get('destNombre') or '').strip():
        return error('Falta el nombre o razón social del destinatario.')
    if not (d.get('remNombre') or '').strip():
        return error('Falta el nombre del remitente (el cliente).')
    if not (d.get('tclNumero') or '').strip():
        return error('Falta el número del TCL previo no contestado.')
    regimen = (d.get('regimen') or 'lct').strip().lower()
    if regimen == 'construccion':
        if not (d.get('categoria') or '').strip():
            return error('Falta la categoría (régimen de la construcción).')
        if not (d.get('jornada') or '').strip():
            return error('Falta la jornada de trabajo (régimen de la construcción).')

    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return error('El servidor no tiene configurada la clave de la API de Claude.', 500)

    try:
        client = anthropic.Anthropic(api_key=api_key)
        relato = ''
        if (d.get('relato') or '').strip():
            relato = redactar_con_ia(client, SYSTEM_PROMPT_TG_INJURIA_DESPIDO, d['relato'])
        jornada_redactada = d.get('jornada', '')
        if regimen == 'construccion' and (d.get('jornada') or '').strip():
            jornada_redactada = redactar_con_ia(
                client, SYSTEM_PROMPT_TG_JORNADA_DESPIDO, d['jornada'])
    except Exception as e:
        print("Error redactando telegrama de despido indirecto con IA:", str(e))
        return error('No se pudo generar el telegrama. Intentá de nuevo en un momento.', 500)

    d2 = dict(d)
    d2['jornada'] = jornada_redactada
    texto = armar_cuerpo_despido_indirecto(d2, relato)
    resp = jsonify({"status": "ok", "texto": texto})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp, 200


# ─── Demanda ART — Rechazo de enfermedad profesional (trámite abreviado) ───
# El escenario que cubre este tipo de escrito es siempre el mismo: la
# Comisión Médica rechazó la contingencia y por eso corresponde el trámite
# abreviado del inciso k. Como el propio "tipo de escrito" ya define ese
# escenario, los bloques de trámite abreviado e impugnación del dictamen van
# siempre fijos, sin checkbox. Los montos y la tabla RIPTE los calcula el
# frontend en vivo (para que se vea cada paso) y acá solo se formatean — no
# se recalculan del lado del servidor, así hay un único lugar con la
# fórmula y el estudio revisa los números antes de generar.

def _numero_a_letras_grupo(n):
    unidades = ['', 'uno', 'dos', 'tres', 'cuatro', 'cinco', 'seis', 'siete', 'ocho', 'nueve']
    especiales = {
        10: 'diez', 11: 'once', 12: 'doce', 13: 'trece', 14: 'catorce', 15: 'quince',
        16: 'dieciséis', 17: 'diecisiete', 18: 'dieciocho', 19: 'diecinueve', 20: 'veinte',
        21: 'veintiuno', 22: 'veintidós', 23: 'veintitrés', 24: 'veinticuatro',
        25: 'veinticinco', 26: 'veintiséis', 27: 'veintisiete', 28: 'veintiocho', 29: 'veintinueve',
    }
    decenas = {30: 'treinta', 40: 'cuarenta', 50: 'cincuenta', 60: 'sesenta',
               70: 'setenta', 80: 'ochenta', 90: 'noventa'}
    centenas = {100: 'cien', 200: 'doscientos', 300: 'trescientos', 400: 'cuatrocientos',
                500: 'quinientos', 600: 'seiscientos', 700: 'setecientos',
                800: 'ochocientos', 900: 'novecientos'}
    if n == 0:
        return ''
    if n < 10:
        return unidades[n]
    if n in especiales:
        return especiales[n]
    if n < 100:
        d = (n // 10) * 10
        u = n % 10
        return decenas[d] + (' y ' + unidades[u] if u else '')
    if n == 100:
        return 'cien'
    c = (n // 100) * 100
    resto = n % 100
    base = centenas[c]
    return base + (' ' + _numero_a_letras_grupo(resto) if resto else '')

def numero_a_letras(n):
    n = int(n)
    if n == 0:
        return 'cero'
    partes = []
    millones = n // 1000000
    resto = n % 1000000
    miles = resto // 1000
    resto2 = resto % 1000
    if millones:
        partes.append('un millón' if millones == 1 else _numero_a_letras_grupo(millones) + ' millones')
    if miles:
        partes.append('mil' if miles == 1 else _numero_a_letras_grupo(miles) + ' mil')
    if resto2:
        partes.append(_numero_a_letras_grupo(resto2))
    return ' '.join(partes)

def monto_en_letras(monto):
    monto = float(monto or 0)
    entero = int(monto)
    centavos = round((monto - entero) * 100)
    if centavos >= 100:
        entero += 1
        centavos = 0
    return 'PESOS %s CON %02d/100' % (numero_a_letras(entero).upper(), centavos)


# En las demandas el patrocinio se cita distinto que en la denuncia: sin CIDI
# ni datos de contacto, solo nombre y matricula. Por defecto van los dos, que
# es como salieron todas las demandas del estudio hasta ahora.
BLOQUES_PATROCINIO_DEMANDA = {
    'ciardiello': 'con el patrocinio letrado de la Dra. Ciardiello Ma. Paula M.P. 1-41227',
    'casih': 'con el patrocinio letrado del Dr. Casih, Pablo Antonio M.P. 1-31142',
    'ambos': ('con el patrocinio letrado de los Dres. Ciardiello Ma. Paula M.P. 1-41227 '
              'y Casih, Pablo Antonio M.P. 1-31142'),
}


def armar_encabezado_actor_art(d):
    nombre = (d.get('actNombre') or '').strip()
    dni = (d.get('actDni') or '').strip()
    fem_act = (d.get('actGenero') or '').strip().lower() == 'femenino'
    nacionalidad = (d.get('actNacionalidad') or ('argentina' if fem_act else 'argentino')).strip()
    estado_civil = (d.get('actEstadoCivil') or '').strip()
    situacion = (d.get('actSituacionLaboral') or '').strip()
    fecha_nac = (d.get('actFechaNac') or '').strip()
    edad = (d.get('actEdad') or '').strip()
    domicilio = (d.get('actDomicilio') or '').strip()

    texto = '%s, DNI %s, %s' % (nombre, dni, nacionalidad)
    if estado_civil:
        texto += ', %s' % estado_civil
    if situacion:
        texto += ', %s' % situacion
    if fecha_nac:
        texto += ', %s el día %s' % ('nacida' if fem_act else 'nacido', fecha_nac)
    if edad:
        texto += ', de %s años de edad' % edad
    bloque_pat = BLOQUES_PATROCINIO_DEMANDA.get(
        d.get('representacion'), BLOQUES_PATROCINIO_DEMANDA['ambos'])
    texto += (
        ', con domicilio real en %s y constituyéndolo a los efectos procesales en calle '
        'Arturo M. Bas N°389 piso 1, oficina "A", ambos de esta ciudad de Córdoba, %s '
        'ante VS comparezco y digo:'
    ) % (domicilio, bloque_pat)
    return texto


def armar_objeto_art(d):
    medico_tit = 'la Dra.' if (d.get('medicoGenero') or '').strip().lower() == 'femenino' else 'el Dr.'
    monto = float(d.get('montoFinal') or 0)
    monto_fmt = '{:,.2f}'.format(monto).replace(',', '_').replace('.', ',').replace('_', '.')

    return (
        'Vengo en tiempo y forma, a iniciar formal demanda ordinaria por resarcimiento de '
        'daños causados a mi persona a causa de las “enfermedades profesionales” contraída '
        'por la prestación de tareas a la orden de mi empleadora – a la que me referiré '
        'infra –, en contra de: %s CUIT: %s, con domicilio en %s, compañía ante la cual, mi '
        'empleadora tenía asegurada su responsabilidad por riesgos del trabajo al momento '
        'del hecho o primera manifestación invalidante (PMI).\n\n'
        'Que, mediante la presente, persigo que, previo los trámites de ley, se determine '
        'el grado y carácter de la incapacidad laborativa que padezco a consecuencia de las '
        'enfermedades profesionales que me aqueja por causa directa e inmediata de la '
        'ejecución de mi trabajo y, en su mérito, se condene a la accionada al pago de las '
        'prestaciones dinerarias previstas en el plexo normativo que se menciona y por la '
        'suma que se reclama de $%s (%s), con más lo que resulte de la prueba a rendirse en '
        'autos, con más intereses, costas y RIPTE, todo en base a los hechos y el derecho '
        'que expongo a continuación.\n\n'
        'La accionada, %s CUIT: %s, se encuentra legitimada pasivamente en virtud del '
        'contrato de afiliación que mantuvo con mi empleadora, %s CUIT %s al momento de la '
        'PMI. Todo ello, de acuerdo con lo prescripto por los arts. 3 inc. 3), 7 y 27 de la '
        'Ley N° 24.557 y demás normativa complementaria, entre ellos se resalta el decreto '
        '1278/00.\n\n'
        'La suma reclamada tiene su origen en que, según los fundamentos de hecho y '
        'derecho, médicos y técnicos que seguidamente se exponen, en la actualidad padezco '
        'un cuadro nosológico, determinante de una incapacidad laboral, provocada por las '
        '(y en relación de causalidad suficiente) ENFERMEDADES PROFESIONALES en los '
        'términos de las Leyes Nº 24.557, Nº 26.773 y N° 27.348. Destaco que recién tomé '
        'conocimiento de la gravedad y alcance de estos una vez otorgado el certificado '
        'médico, de fecha %s, extendido por %s %s Especialista en medicina del trabajo M.P. '
        '%s que se adjunta a la presente, siendo dicha fecha, de extensión del certificado, '
        'la PMI en los términos de ley. Ahora bien, por el grado e importancia de dicha '
        'enfermedad, la incapacidad que en consecuencia sobreviene y el derecho que a '
        'continuación expongo solicito que, previo los trámites de ley, se determine el '
        'grado y carácter de la incapacidad laborativa que padezco realmente por causa '
        'directa e inmediata de la ejecución de mi trabajo y, en su mérito, se condene a la '
        'accionada al pago de las prestaciones dinerarias previstas por la suma mencionada '
        'ut supra.'
    ) % (
        d.get('artNombre', ''), d.get('artCuit', ''), d.get('artDomicilio', ''),
        monto_fmt, monto_en_letras(monto),
        d.get('artNombre', ''), d.get('artCuit', ''),
        d.get('empNombre', ''), d.get('empCuit', ''),
        d.get('fechaCert', ''), medico_tit, d.get('medicoNombre', ''), d.get('medicoMp', ''),
    )


BLOQUE_TRAMITE_ABREVIADO_ART = (
    'En base al análisis relatado ut supra, solicito a V.S. se imprima el trámite de '
    'juicio abreviado a la presente demanda, ya que resulta evidente que se encuentran '
    'cumplimentados los requisitos previstos por el art. 83 bis de la Ley N° 7987 (texto '
    'según Ley N° 10.596, vigente desde el 01/04/2021) para la procedencia de la vía '
    'abreviada que se ha impetrado.\n\n'
    'En agosto del año 2025 se implementó la vigencia del inciso k que nos dice:\n\n'
    'Art 83 bis: (…) k) Demandas derivadas del Régimen de Riesgos del Trabajo por '
    'accidentes de trabajo o enfermedades profesionales cuya contingencia, hecho '
    'generador, relación causal o calificación médico legal haya sido rechazada por la '
    'Comisión Médica Jurisdiccional dependiente de la Superintendencia de Riesgos del '
    'Trabajo, y l) (…).”\n\n'
    'Estamos frente a una hipótesis de dicho inciso, encontrándose todos los recaudos '
    'exigidos por la ley.'
)


def armar_hechos_antecedentes_art(d, tareas_redactadas):
    partes = []
    partes.append(
        'Pongo de manifiesto que, me encontraba trabajando en relación de dependencia '
        'laboral, económica, jurídica y técnica para %s, con domicilio social en %s. '
        'Contaba con una antigüedad reconocida desde el %s haciéndolo de manera continua e '
        'ininterrumpida hasta el día %s, fecha en que se extinguió la relación laboral.\n\n'
        'Que ingresé a trabajar en perfectas condiciones físicas y sin mengua de mi '
        'capacidad laborativa, cumpliendo una jornada laboral %s, todo ello conforme '
        'modalidad impuesta por mi empleadora.\n\n'
        'Que cumplía tareas estando jerarquizado, según consta en mis recibos de haberes, '
        'dentro de la categoría “%s” conforme CCT que regula la actividad de la patronal.'
        % (d.get('empNombre', ''), d.get('empDomicilio', ''),
           d.get('fechaIngreso', ''), d.get('fechaEgreso', ''),
           d.get('jornada', ''), d.get('categoria', ''))
    )
    if tareas_redactadas:
        partes.append(tareas_redactadas)
    return '\n\n'.join(partes)


def armar_sintomas_y_certificado_art(d, sintomas_redactados):
    partes = []
    if sintomas_redactados:
        partes.append(sintomas_redactados)
    medico_tit = 'la Dra.' if (d.get('medicoGenero') or '').strip().lower() == 'femenino' else 'el Dr.'
    partes.append(
        'Por ese motivo consulté con un especialista en medicina del trabajo %s %s MP %s, '
        'quien, con fecha %s tras una revisión, exámenes y estudios médicos, constató y '
        'confirmó la patología que padezco a raíz de mis tareas laborales habituales. '
        'Determinó que ella es consecuencia directa e inmediata de la actividad laboral '
        'que realicé.'
        % (medico_tit, d.get('medicoNombre', ''), d.get('medicoMp', ''), d.get('fechaCert', ''))
    )
    return ' '.join(partes)


def armar_tramite_srt_art(d):
    partes = []
    partes.append(
        'Realicé, mediante telegrama ley 23.789 N° CD %s de fecha %s, la denuncia ante la '
        'demandada de las enfermedades profesionales que padezco.'
        % (d.get('tclNumero', ''), d.get('tclFecha', ''))
    )
    # La ART suele citar mas de una vez por distintos motivos (evaluacion
    # general, un estudio puntual, etc.), asi que van todas las que se hayan
    # cargado, en el orden en que las agregaron.
    for control in (d.get('controles') or []):
        fecha_control = (control.get('fecha') or '').strip()
        if not fecha_control:
            continue
        texto_control = 'Con fecha %s, la demandada me cita a control.' % fecha_control
        detalle = (control.get('detalle') or '').strip()
        if detalle:
            texto_control += ' %s' % detalle
        partes.append(texto_control)
    # No todos los casos tienen un rechazo directo de la ART por CD: a veces
    # la ART informa "sin incapacidad" directamente a la SRT (con el alta
    # medica como fundamento) y el expediente arranca sin que medie una carta
    # documento de rechazo hacia el trabajador.
    hubo_cd_rechazo = (d.get('cdRechazoNumero') or '').strip()
    if hubo_cd_rechazo:
        partes.append(
            'El día %s la demandada remite CD Nº %s, rechazando las patologías denunciadas por '
            'considerar que no me encontraba expuesto a los agentes de riesgo, sin tomar en '
            'consideración la relación de causalidad existente entre mis tareas laborales y '
            'las dolencias que padezco.'
            % (d.get('cdRechazoFecha', ''), d.get('cdRechazoNumero', ''))
        )
    partes.append(
        '%s expediente ante la SRT con fecha %s. El expediente fue tramitado bajo el N°%s.'
        % ('Debido a ello inicio' if hubo_cd_rechazo else 'Se inicia',
           d.get('expteFecha', ''), d.get('expteNumero', ''))
    )
    if (d.get('informeTecnicoFecha') or '').strip():
        partes.append(
            'Con fecha %s se emite el Informe Técnico Médico, en donde decide la SRT no '
            'citar a junta médica a los fines de poder revisar al actor. Lo cual a todas '
            'luces refleja el poco interés y que la negativa ya está predispuesta de '
            'antemano.' % d['informeTecnicoFecha']
        )
    if (d.get('dictamenFecha') or '').strip():
        texto = 'Con fecha %s se emitió Dictamen Médico:' % d['dictamenFecha']
        if (d.get('dictamenTexto') or '').strip():
            texto += '\n\n“%s”' % d['dictamenTexto'].strip()
        partes.append(texto)
    return '\n\n'.join(partes)


BLOQUE_IMPUGNA_DICTAMEN_ART = (
    'El dictamen ha incurrido en error de juzgamiento al omitir o infravalorar '
    'científicamente tanto las patologías que padezco como los estudios médicos '
    'practicados, lo que se demostrará en la etapa procesal oportuna y se acredita '
    'científicamente en este acto con el certificado médico que se acompaña, y que, de '
    'haber sido tenido en cuenta, se hubiera determinado adecuadamente mi incapacidad. El '
    'dictamen no ha tenido en cuenta la totalidad de los elementos ofrecidos como prueba '
    'ni han expuesto los fundamentos científicos mínimos, respetando así las reglas de la '
    'experiencia de las que debe valerse todo dictamen. Esto permite concluir que estamos '
    'frente a un dictamen que adolece del vicio de falta de fundamentación científica y '
    'que llevan la tacha de arbitrarios, por ende, adolecen de un vicio que acarrea su '
    'nulidad absoluta e insalvable.\n\n'
    'La rectificación del dictamen o su anulación y reemplazo por la pericia médica en '
    'autos, me permitirán el acceso a una indemnización acorde a mi minusvalía.'
)


def armar_subsidiariamente_recurre_art(d):
    monto = float(d.get('montoFinal') or 0)
    monto_fmt = '{:,.2f}'.format(monto).replace(',', '_').replace('.', ',').replace('_', '.')
    comision = (d.get('comisionMedicaNumero') or '').strip()
    return (
        'Sin perjuicio de lo expuesto anteriormente, y de manera subsidiaria vengo a '
        'recurrir el arbitrario Dictamen emitido por la Comisión Médica Nº %s, solicitando '
        'que el Tribunal oportunamente lo declare nulo y condene a %s CUIT: %s con '
        'domicilio real en %s, a abonar la prestación dineraria correspondiente a la '
        'incapacidad parcial, permanente y definitiva generada por la enfermedad '
        'profesional que sufro, prevista en la LRT por la suma de $%s (%s), y/o lo que en '
        'más o en menos surja de la prueba a rendirse, con más intereses y costas.\n\n'
        'Como he dicho anteriormente, mi reclamo se inició mediante la denuncia a la '
        'demandada de la afección que me aqueja, posteriormente recurrí a la Comisión '
        'Médica Nº %s, solicitando su intervención. La mencionada Comisión Médica, se '
        'expidió determinando la existencia de una enfermedad de carácter laboral, pero '
        'infravalorando la incapacidad que esta me genera.\n\n'
        'En relación a los agravios que me ocasiona el mencionado pronunciamiento, me '
        'remito a lo expresado ut-supra bajo el epígrafe: “IMPUGNA DICTAMEN COMISIÓN '
        'MÉDICA - EXPRESA AGRAVIOS”.'
        % (comision, d.get('artNombre', ''), d.get('artCuit', ''), d.get('artDomicilio', ''),
           monto_fmt, monto_en_letras(monto), comision)
    )


def _fmt_monto(v):
    try:
        return '{:,.2f}'.format(float(v or 0)).replace(',', '_').replace('.', ',').replace('_', '.')
    except (TypeError, ValueError):
        return '0,00'


def armar_rubros_reclamados_art(d):
    partes = []
    partes.append(
        'El siguiente cálculo será realizado de acuerdo a lo establecido por las Leyes '
        '24.557, 26.773, 27.348 y el Decreto 1694/2009.'
    )
    partes.append('PMI: %s\nÍndice base mes de la PMI: %s' % (
        d.get('fechaPmi', ''), _fmt_monto(d.get('ripteDestino'))))
    partes.append('[[TABLA_RIPTE]]')
    partes.append(
        'Valor mensual de ingreso base: $%s / 12: $%s'
        % (_fmt_monto(d.get('ripteTotal')), _fmt_monto(d.get('vibm')))
    )
    partes.append(
        'Edad a la fecha de PMI: 65 / %s: %s'
        % (d.get('edadPmi', ''), d.get('coefEdad', ''))
    )
    partes.append(
        'Cálculo indemnización: 53 x $%s x %s x %s %%: $%s'
        % (_fmt_monto(d.get('vibm')), d.get('coefEdad', ''),
           d.get('porcentajeIncapacidad', ''), _fmt_monto(d.get('indemnizacionBase')))
    )
    if (d.get('pisoResolucion') or '').strip():
        corresponde = 'SI' if d.get('pisoCorresponde') else 'NO'
        partes.append(
            '•PISO MÍNIMO %s: $ %s * %s%%: $ %s (%s corresponde su aplicación por ser %s '
            'a la Fórmula).'
            % (d['pisoResolucion'], _fmt_monto(d.get('pisoMonto')),
               d.get('porcentajeIncapacidad', ''), _fmt_monto(d.get('pisoAplicable')),
               corresponde, 'superior' if d.get('pisoCorresponde') else 'inferior')
        )
    partes.append(
        'A los montos parciales detallados anteriormente se le debe adicionar la '
        'compensación del 20%% dispuesta por el art. 3 de la Ley 26.773 por los daños '
        'efectivamente sufridos por el trabajador y que palmariamente no son cubiertos por '
        'los métodos de cálculo establecidos por la Ley 24.557.\n\n'
        'En virtud de lo manifestado previamente, solicito se adicione dicho 20%%, de '
        'acuerdo al siguiente cálculo:\n'
        '1) $%s\n'
        '2) $%s + 20%% art 3 ley 26773: $%s\n'
        '3) $%s + $%s: $%s'
        % (_fmt_monto(d.get('indemnizacionBase')),
           _fmt_monto(d.get('indemnizacionBase')), _fmt_monto(d.get('monto20')),
           _fmt_monto(d.get('indemnizacionBase')), _fmt_monto(d.get('monto20')),
           _fmt_monto(d.get('montoFinal')))
    )
    partes.append(
        'El monto total que me corresponde en concepto de prestación dineraria por la '
        'incapacidad que padezco, a causa de las patologías que sufro como consecuencia de '
        'las tareas prestadas para con mi empleador, asciende a la suma de $%s (%s).'
        % (_fmt_monto(d.get('montoFinal')), monto_en_letras(d.get('montoFinal')))
    )
    partes.append(
        'El monto reclamado es estimativo y al solo efecto de cubrir recaudos procesales, '
        'quedando el monto definitivo a lo que en más o en menos resulte de la prueba a '
        'rendirse en autos, todo ello con más actualizaciones, intereses, costas que el '
        'Tribunal de mérito estime correcta, desde la fecha de PMI y hasta la fecha del '
        'efectivo pago.\n\n'
        'Dejo expresamente impugnada la aplicación de la Resolución Nº 1039/2019 y 332/23 '
        'en razón de que la misma atenta con el Principio de Reparación suficiente de los '
        'siniestros/enfermedades profesionales, derecho de propiedad y seguridad jurídica.'
    )
    partes.append(
        'La sumatoria simple de los índices no refleja verdaderamente el paso del tiempo y '
        'como afecta la inflación a las remuneraciones percibidas por los trabajadores '
        'activos. También implicaría aceptar un detrimento confiscatorio del crédito del '
        'trabajador, sujeto de especial protección al decir de la CSJN, más aún con un '
        'grado de incapacidad que limita su reinserción laboral. Aplicar dichas '
        'resoluciones quiebra la jerarquía y coherencia normativa impuesta por la propia '
        'Constitución.'
    )
    partes.append(
        'En razón de ello, dejo planteado que se aplique al caso concreto la normativa que '
        'sea la más favorable para el trabajador conforme lo establece la LCT. Y solicito '
        'se aplique un interés compensatorio del 11% anual para contrarrestar los efectos '
        'de la inflación y de las demoras que tienen los procesos judiciales. Subsidiariamente '
        'solicito aplicación del art. 1109 del CCIV y Art. 1113 CCIV. Siempre lo que sea más '
        'favorable al Trabajador.'
    )
    return '\n\n'.join(partes)


def armar_prueba_art(d):
    partes = []
    partes.append(
        'a.-) CONFESIONAL: Del representante legal de %s CUIT: %s con facultades para '
        'absolver posiciones de acuerdo a los estatutos, en forma personal e indelegable '
        'quien deberá ser citado a la vista de la causa bajo apercibimientos de ley, para '
        'que deponga a tenor del pliego que en esa oportunidad se presentará.'
        % (d.get('artNombre', ''), d.get('artCuit', ''))
    )

    testigos = d.get('testigos') or []
    if testigos:
        lineas = ['- %s con domicilio en %s;' % (t.get('nombre', ''), t.get('domicilio', ''))
                  for t in testigos if (t.get('nombre') or '').strip()]
        partes.append(
            'b.-) TESTIMONIAL\nSolicito se fije fecha y hora de audiencia para receptar la '
            'declaración testimonial de las siguientes personas:\n' + '\n'.join(lineas)
        )

    doc_items = [it.strip() for it in (d.get('documentalItems') or []) if (it or '').strip()]
    if (d.get('documentalOtros') or '').strip():
        doc_items.append(d['documentalOtros'].strip())
    if doc_items:
        letras = 'abcdefghijklmnopqrstuvwxyz'
        lista = ' '.join('%s.-) %s' % (letras[i], it) for i, it in enumerate(doc_items))
        partes.append(
            'c.-) DOCUMENTAL – INSTRUMENTAL: Acompaña en adjunto: %s\n\nLa documental se '
            'adjunta en formato PDF, declarando que es copia fiel de la original, poniendo '
            'a disposición para cuando el Tribunal lo requiera.' % lista
        )

    partes.append(
        'd.-) EXHIBICIÓN: Solicito día y hora de audiencia a los fines de la exhibición '
        'por parte de la demandada de la documental que se detalla a continuación '
        'agregando copia certificada en los autos del rubro: a) Examen de ingreso '
        '(Pre-ocupacional) y Exámenes médicos periódicos que debieron realizar %s la '
        'empleadora en base a lo dispuesto por el Decreto 170/96 y resolución SRT 43/97 '
        'Art. 3, Anexo II y Res. 28/98. b) Plan de mejoras, constancias de visitas y '
        'fiscalización del efectivo cumplimiento de las medidas de higiene y seguridad que '
        'debió realizar la ART al empleador %s. c) Para el supuesto caso que la ART haya '
        'hecho denuncia en la Superintendencia del Riesgo del Trabajo por incumplimiento '
        'de la Empresa de los puntos a) y b) de este apartado, constancia certificada de '
        'dicha DENUNCIA. d.-) Constancia de capacitación al actor teniendo en cuenta a los '
        'riesgos que se encuentra expuesto por su actividad. e-) Copia certificada del '
        'contrato de cobertura suscripto entre %s y %s, indicando fecha de inicio y '
        'finalización, si éste hubiese caducado. f-) Denuncia de Enfermedad laboral '
        'sufrida por el actor, informe y constancias de atenciones médicas brindadas al '
        'actor.'
        % (d.get('artNombre', ''), d.get('empNombre', ''), d.get('artNombre', ''), d.get('empNombre', ''))
    )

    texto_reconocimiento = (
        'e.-) RECONOCIMIENTO\nSolicito se fije día y hora de audiencia a los fines del '
        'reconocimiento por parte de la demandada: a.-) Recepción, autenticidad, cuerpo y '
        'contenido del TCL CD Nº %s ofrecida en el punto c.-) de la DOCUMENTAL.'
        % d.get('tclNumero', '')
    )
    if (d.get('cdRechazoNumero') or '').strip():
        texto_reconocimiento += (
            ' b.-) Emisión, autenticidad, cuerpo y contenido, de la CD Nº %s ofrecida en el '
            'mismo punto de la DOCUMENTAL.' % d['cdRechazoNumero']
        )
    partes.append(texto_reconocimiento)

    perito_control = ''
    if (d.get('peritoControlNombre') or '').strip():
        perito_control = (
            '\n\nPERITOS DE CONTROL: Se deja ofrecido como perito de control médico a %s '
            'M.P. %s con domicilio en %s.'
            % (d['peritoControlNombre'], d.get('peritoControlMp', ''), d.get('peritoControlDomicilio', ''))
        )
    partes.append(
        'f.-) DICTÁMENES PERICIALES:\n\n'
        'a) PERICIAL MEDICA: Solicito se fije día y hora de audiencia a los fines de un '
        'sorteo de un perito médico especialista en Medicina del Trabajo, a fin de que '
        'presente un informe fundado, previa revisación del actor y de los estudios '
        'médicos que ésta acompañe, manifestando en base a lo establecido por la Ley '
        '24.557, y dto. 658/96 si; 1) El actor padece las dolencias denunciadas. 2) Si las '
        'mismas le producen incapacidad, y en su caso si la misma es de carácter '
        'permanente y definitiva. 3) Si las patologías son consecuencia de la actividad '
        'laboral que desarrollaba el actor. Fundamentando adecuadamente la relación de '
        'causalidad de las tareas que cumplía y la patología que éste padece. 4) Cualquier '
        'otra consideración que el perito crea oportuno realizar, en base a su experiencia '
        'y pericias anteriores. 5) Calificación médico legal, si las patología '
        'incapacitantes deben calificarse como enfermedad profesional. 6) En su caso '
        'establecer grado de incapacidad, teniendo en cuenta los factores de ponderación '
        'dto. 659/96.\n\n'
        'Se solicita expresamente se respondan a todas y cada una de las preguntas '
        'formuladas, dando fundamento de las respuestas, bajo apercibimiento de solicitar '
        'la nulidad del acto pericial.%s\n\n'
        'b.-) PERICIA TECNICA: Solicito se fije día y hora de audiencia a los fines de que '
        'se designe Perito Oficial Ingeniero en Higiene y Seguridad del Trabajo, a los '
        'siguientes fines: 1.) Para que se constituya en la sede de %s, ubicado en %s para '
        'recabar todo tipo de información relacionada a los riesgos inherentes a los que '
        'estuvo expuesto el actor y las causas que pudieron ocasionar la enfermedad '
        'profesional denunciada en autos. 2.) Indique qué consecuencias produciría en una '
        'persona estar sometida a las tareas descriptas por la actora en la demanda. 3.) '
        'Indique si las tareas que realizó la actora pueden ser causa de la enfermedad '
        'denunciada en esta demanda. 4.) Aplicación por parte de la demandada de las '
        'medidas de higiene y seguridad previstas en la Ley Nº 19.587 y sus '
        'reglamentaciones (Decretos Nº 351/79, 991/96, 1338/96, 617/97, 249/07 y/o '
        'cualquier otra normativa de aplicación), a las tareas que efectuaba el actor en '
        'su carácter de dependiente de %s. 5.) Aplicación por parte de la demandada de las '
        'medidas de higiene y seguridad previstas en la Ley Nº 24.557, sus '
        'reglamentaciones (Decretos Nº 170/96, 334/96, 658/96, 1278/00) y las resoluciones '
        'emitidas por la Superintendencia de Riesgos del Trabajo, que debía adoptar en el '
        'puesto de trabajo denunciado por la actora en la demanda; 6.) Informe si en '
        'relación a las tareas que denuncia el actor en su demanda, y en especial teniendo '
        'en cuenta las patologías que la accionante padece según el certificado médico que '
        'se adjunta a la demanda, se adoptaron por parte de la demandada las medidas de '
        'higiene y seguridad y prevención de riesgos del trabajo que prevé la '
        'reglamentación vigente; 7.) Informe si el actor, en cumplimiento de las tareas '
        'encomendadas por su empleadora, se encontró expuesto a los agentes de riesgos '
        'previstos en el Decreto 658/1996, en especial los denunciados oportunamente en '
        'escrito de demanda; 8.) Informe si de las constancias documentales emanadas de la '
        'demandada surge la capacitación a los empleados en relación a los riesgos del '
        'trabajo que se encuentran expuestos y si se proveyó a los trabajadores de los '
        'elementos de protección adecuados para evitar o impedir accidentes de trabajo '
        'y/o las enfermedades profesionales relatadas en la demanda.\n\n'
        'Toda la información arriba requerida deberá ser de fecha anterior a la denuncia '
        'efectuada por el actor ante la demandada.\n\n'
        'Formulo reserva de ampliar puntos de pericia y de designar perito de control en '
        'los términos de los arts. 262 y 263 del CPCC.'
        % (perito_control, d.get('empNombre', ''), d.get('empDomicilio', ''), d.get('empNombre', ''))
    )

    salud = (d.get('saludEstablecimiento') or '').strip()
    partes.append(
        'g.-) INFORMATIVA: a) Al ARCA, a los fines que informe remuneraciones declaradas '
        'por %s con referencia al actor. b.-) A la COMISION MEDICA LOCAL, a los fines que '
        'remita copia certificada de todas las actuaciones correspondientes al Expediente '
        'SRT N° %s.%s'
        % (d.get('empNombre', ''), d.get('expteNumero', ''),
           (' c.-) A %s: a los fines de que remita al Tribunal original y/o copia '
            'certificada de la Historia Clínica a nombre del actor.' % salud) if salud else '')
    )

    return '\n\n'.join(partes)


BLOQUE_CIERRE_ART = (
    'VIII.-) PRINCIPIO “IURA NOVIT CURIA”:\n\n'
    'En el caso de que existiera un error en el encuadramiento de los hechos en la ley, '
    'pido en virtud del mencionado principio, que el Juzgador corrija cualquier cita '
    'equivocada, subsumiendo la norma legal al marco fáctico citado.\n\n'
    'IX.-) RESERVA DEL CASO FEDERAL:\n\n'
    'Formulo reserva del caso federal y del recurso extraordinario que autoriza el art. 14 '
    'de la Ley 48 y doctrina de la arbitrariedad. Lo dicho es aplicable para el supuesto '
    'que se dicte una resolución contraria a mis legítimos intereses, por cuanto en tal '
    'supuesto nos encontraríamos ante la violación de derechos de raigambre '
    'constitucional, tales como el de propiedad, igualdad ante la ley, defensa en juicio, '
    'entre otros, tutelados por nuestra Ley Fundamental.\n\n'
    'X.-) DERECHO:\n\n'
    'Fundamento mi pretensión en los siguientes Arts. 75 y c.c. de la L.C.T., 6, 8 apartado '
    '1, art 14 apartado 2 inc. A y c.c. de la Ley de Riesgos de Trabajo y los Arts. 14 bis, '
    '16, 75 inc. 12, 116 y c.c. de la Constitución Nacional; los Tratados Internacionales a '
    'los que nuestro país adhirió; los Arts. 1, 9, 10 L.P.T. 7987; Ley 10596, y las más '
    'moderna doctrina y jurisprudencia aplicable al presente caso.\n\n'
    'PETITUM:\n\n'
    'Por todo lo expuesto a V.S. solicito: a.-) Me tenga por presentado, por parte y con '
    'el domicilio constituido. b.-) Admita la presente demanda. c.-) Se sirva correr los '
    'plazos para contestar la demanda y/o expresar agravios. d.-) Tenga por ofrecida la '
    'prueba. e.-) Tenga presente la reserva del caso federal. f.-) Tenga por acompañada la '
    'documental mencionada.\n\n'
    'SERA JUSTICIA.'
)


def armar_cuerpo_demanda_art(d, tareas_redactadas, sintomas_redactados):
    partes = []
    partes.append(armar_encabezado_actor_art(d))
    partes.append('I.-) OBJETO:\n\n' + armar_objeto_art(d))
    partes.append('II.-) SOLICITA TRÁMITE ABREVIADO (PDA):\n\n' + BLOQUE_TRAMITE_ABREVIADO_ART)
    partes.append(
        'III.-) HECHOS ANTECEDENTES DE LA RELACIÓN LABORAL:\n\n1º) '
        + armar_hechos_antecedentes_art(d, tareas_redactadas)
        + '\n\n2º) ' + armar_sintomas_y_certificado_art(d, sintomas_redactados)
        + '\n\n3º) ' + armar_tramite_srt_art(d)
    )
    partes.append('IV.-) IMPUGNA DICTAMEN COMISIÓN MÉDICA- EXPRESA AGRAVIOS:\n\n' + BLOQUE_IMPUGNA_DICTAMEN_ART)
    partes.append('V.-) SUBSIDIARIAMENTE RECURRE DICTAMEN DE COMISIÓN MÉDICA:\n\n' + armar_subsidiariamente_recurre_art(d))
    partes.append('VI.-) RUBROS RECLAMADOS:\n\n' + armar_rubros_reclamados_art(d))
    partes.append('VII.-) OFRECE PRUEBA:\n\n' + armar_prueba_art(d))
    partes.append(BLOQUE_CIERRE_ART)
    return '\n\n'.join(p for p in partes if p and p.strip())


SYSTEM_PROMPT_TAREAS_ART = (
    "Sos un asistente de redacción para un estudio jurídico laboralista de Córdoba, "
    "Argentina. Vas a recibir una descripción informal de las tareas físicas que hacía un "
    "trabajador (movimientos, posturas, pesos, repetición) y tenés que convertirla en uno "
    "o más párrafos formales, en primera persona, para una demanda por enfermedad "
    "profesional, describiendo en detalle el esfuerzo físico y las condiciones "
    "ergonómicas de la tarea, arrancando de forma natural (por ejemplo: 'Que mis tareas "
    "consistían en...').\n\n"
    "Reglas:\n"
    "- Nunca inventes datos, pesos ni cantidades que no estén en el texto que te pasaron.\n"
    "- Devolvé únicamente el/los párrafo(s) final(es), sin comentarios antes o después."
)

SYSTEM_PROMPT_SINTOMAS_ART = (
    "Sos un asistente de redacción para un estudio jurídico laboralista de Córdoba, "
    "Argentina. Vas a recibir una descripción informal de los síntomas o dolencias que "
    "empezó a sentir un trabajador a raíz de sus tareas, y tenés que convertirla en una "
    "frase corrida, en primera persona, que continúe naturalmente la oración: 'Que a "
    "razón de las actividades realizadas durante la ejecución de mis labores, "
    "[TU TEXTO].'\n\n"
    "Reglas:\n"
    "- Arrancá en minúscula, sin punto final si es posible integrarlo, ya que tu frase "
    "continúa una oración que ya empezó.\n"
    "- Nunca inventes síntomas ni datos que no estén en el texto que te pasaron.\n"
    "- Devolvé únicamente la frase, sin comentarios antes o después."
)


@app.route('/generar-demanda-art', methods=['POST', 'OPTIONS'])
def generar_demanda_art():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    d = request.get_json(silent=True) or {}

    def error(msg, code=400):
        resp = jsonify({"status": "error", "detalle": msg})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, code

    if not (d.get('actNombre') or '').strip():
        return error('Falta el nombre del actor.')
    if not (d.get('artNombre') or '').strip():
        return error('Falta el nombre de la ART demandada.')
    if not (d.get('montoFinal') or 0):
        return error('Falta el monto reclamado (completá la calculadora RIPTE).')

    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return error('El servidor no tiene configurada la clave de la API de Claude.', 500)

    try:
        client = anthropic.Anthropic(api_key=api_key)
        tareas = ''
        if (d.get('tareasRelato') or '').strip():
            tareas = redactar_con_ia(client, SYSTEM_PROMPT_TAREAS_ART, d['tareasRelato'])
        sintomas = ''
        if (d.get('sintomasRelato') or '').strip():
            sintomas = redactar_con_ia(client, SYSTEM_PROMPT_SINTOMAS_ART, d['sintomasRelato'])
    except Exception as e:
        print("Error redactando demanda ART con IA:", str(e))
        return error('No se pudo generar la demanda. Intentá de nuevo en un momento.', 500)

    texto = armar_cuerpo_demanda_art(d, tareas, sintomas)
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    resp = jsonify({"status": "ok", "texto": texto})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp, 200


def construir_docx_demanda_art(texto, actor, filas_ripte):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    estilo = doc.styles['Normal']
    estilo.font.name = 'Times New Roman'
    estilo.font.size = Pt(12)
    estilo.paragraph_format.line_spacing = 1.15
    estilo.paragraph_format.space_after = Pt(0)

    titulos = (
        'I.-) OBJETO', 'II.-) SOLICITA', 'III.-) HECHOS', 'IV.-) IMPUGNA',
        'V.-) SUBSIDIARIAMENTE', 'VI.-) RUBROS', 'VII.-) OFRECE PRUEBA',
        'VIII.-) PRINCIPIO', 'IX.-) RESERVA', 'X.-) DERECHO', 'PETITUM',
    )

    bloques = [b.strip('\n') for b in texto.split('\n\n') if b.strip()]
    for bloque in bloques:
        if bloque.strip() == '[[TABLA_RIPTE]]':
            tabla = doc.add_table(rows=1, cols=6)
            tabla.style = 'Table Grid'
            cab = tabla.rows[0].cells
            for i, txt in enumerate(['Mes', 'Remuneración', 'RIPTE origen', 'RIPTE destino', 'Coeficiente', 'Total']):
                cab[i].text = txt
                cab[i].paragraphs[0].runs[0].bold = True
            for fila in filas_ripte:
                celdas = tabla.add_row().cells
                celdas[0].text = str(fila.get('mes', ''))
                celdas[1].text = '$ ' + _fmt_monto(fila.get('remuneracion'))
                celdas[2].text = str(fila.get('ripteOrigen', ''))
                celdas[3].text = str(fila.get('ripteDestino', ''))
                celdas[4].text = str(fila.get('coeficiente', ''))
                celdas[5].text = '$ ' + _fmt_monto(fila.get('total'))
            doc.add_paragraph()
            continue

        if bloque.startswith(titulos):
            p = doc.add_paragraph()
            p.add_run(bloque).bold = True
            continue

        if bloque.startswith('“') and bloque.endswith('”'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.right_indent = Cm(1.5)
            p.add_run(bloque).italic = True
            continue

        p = doc.add_paragraph(bloque)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_paragraph()
    doc.add_paragraph()
    abogados = FIRMAS_ABOGADOS.get(actor.get('representacion'), FIRMAS_ABOGADOS['ambos'])
    tabla_firma = doc.add_table(rows=1, cols=1 + len(abogados))
    celdas = tabla_firma.rows[0].cells
    _agregar_firma(celdas[0], (actor.get('nombre') or '').strip(),
                    ('DNI %s' % actor['dni']) if actor.get('dni') else '')
    for i, (nombre_ab, titulo_ab, mp_ab) in enumerate(abogados):
        _agregar_firma(celdas[1 + i], nombre_ab, titulo_ab, mp_ab)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.route('/generar-demanda-art-docx', methods=['POST', 'OPTIONS'])
def generar_demanda_art_docx():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    body = request.get_json(silent=True) or {}
    texto = (body.get('texto') or '').strip()
    actor = body.get('actor') or {}
    filas_ripte = body.get('filasRipte') or []

    if not texto:
        resp = jsonify({"status": "error", "detalle": "Falta el texto de la demanda."})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 400

    try:
        buffer = construir_docx_demanda_art(texto, actor, filas_ripte)
    except Exception as e:
        print("Error generando docx de demanda ART:", str(e))
        resp = jsonify({"status": "error", "detalle": "No se pudo generar el archivo Word."})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 500

    nombre_archivo = 'Demanda ART - %s.docx' % ((actor.get('nombre') or 'escrito').strip() or 'escrito')
    resp = send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=nombre_archivo
    )
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp


@app.route('/generar-escrito', methods=['POST', 'OPTIONS'])
def generar_escrito():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    body = request.get_json(silent=True) or {}
    tipo = body.get('tipo')
    cliente = body.get('cliente') or {}
    denunciado = body.get('denunciado') or {}
    relacion = body.get('relacion') or {}
    correspondencia = body.get('correspondencia') or []

    def error(msg, code=400):
        resp = jsonify({"status": "error", "detalle": msg})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, code

    if tipo != 'denuncia_trabajo':
        return error('Tipo de escrito no reconocido.')
    if not cliente.get('nombre'):
        return error('Faltan los datos del cliente.')
    if not denunciado.get('esPersona') and not denunciado.get('esEmpresa'):
        return error('Marcá si el denunciado es persona física, empresa, o ambas.')
    if not (relacion.get('tareas') or '').strip():
        return error('Falta la descripción de las tareas.')
    if not (relacion.get('relato') or '').strip():
        return error('Falta el relato de los hechos.')
    for item in correspondencia:
        if item.get('huboRespuesta') and not (item.get('respuesta') or {}).get('contenido', '').strip():
            return error('Marcaste que hubo respuesta de la empresa en un telegrama pero falta el contenido de esa respuesta.')

    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return error('El servidor no tiene configurada la clave de la API de Claude.', 500)

    try:
        client = anthropic.Anthropic(api_key=api_key)

        parrafo_tareas = redactar_con_ia(
            client, SYSTEM_PROMPT_TAREAS,
            "TEXTO INFORMAL DE LAS TAREAS:\n%s" % relacion.get('tareas', '')
        )
        parrafo_relato = redactar_con_ia(
            client, SYSTEM_PROMPT_RELATO,
            "RELATO INFORMAL DE LOS HECHOS:\n%s" % relacion.get('relato', '')
        )

        for item in correspondencia:
            if item.get('emisor') == 'empresa' and item.get('modo') == 'resumir' and (item.get('contenido') or '').strip():
                item['contenido'] = redactar_con_ia(client, SYSTEM_PROMPT_RESUMEN_RESPUESTA, item['contenido'])
            respuesta = item.get('respuesta')
            if respuesta and respuesta.get('modo') == 'resumir' and (respuesta.get('contenido') or '').strip():
                respuesta['contenido'] = redactar_con_ia(client, SYSTEM_PROMPT_RESUMEN_RESPUESTA, respuesta['contenido'])

        # Si el domicilio vino en mayusculas sostenidas (comun al copiar de un
        # DNI o formulario), se corrige solo la capitalizacion -- nunca las
        # palabras, numeros o el contenido.
        cliente['domicilio'] = normalizar_mayusculas_domicilio(client, cliente.get('domicilio', ''))
        denunciado['domicilio'] = normalizar_mayusculas_domicilio(client, denunciado.get('domicilio', ''))
    except Exception as e:
        print("Error generando parrafos con IA:", str(e))
        return error('No se pudo generar el escrito. Intentá de nuevo en un momento.', 500)

    texto_final = armar_escrito_denuncia_trabajo(cliente, denunciado, relacion, correspondencia, parrafo_tareas, parrafo_relato)

    resp = jsonify({"status": "ok", "texto": texto_final})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp, 200

# Lee un telegrama/CD adjunto (foto o PDF) y extrae fecha, numero de
# seguimiento y contenido. Es una transcripcion asistida: el frontend siempre
# la muestra en un campo editable para que se revise contra el original antes
# de usarla en un escrito.
@app.route('/extraer-telegrama', methods=['POST', 'OPTIONS'])
def extraer_telegrama():
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return resp, 200

    def error(msg, code=400):
        resp = jsonify({"status": "error", "detalle": msg})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, code

    body = request.get_json(silent=True) or {}
    archivo_b64 = body.get('archivo')
    media_type = body.get('mediaType') or 'application/pdf'

    if not archivo_b64:
        return error('Falta el archivo.')

    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return error('El servidor no tiene configurada la clave de la API de Claude.', 500)

    if media_type == 'application/pdf':
        content_block = {"type": "document", "source": {"type": "base64", "media_type": media_type, "data": archivo_b64}}
    elif media_type.startswith('image/'):
        content_block = {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": archivo_b64}}
    else:
        return error('Formato de archivo no soportado. Subí una foto o un PDF.')

    system_prompt = (
        "Sos un asistente que transcribe telegramas y cartas documento (CD) de Correo Argentino "
        "para un estudio jurídico. Te paso una imagen o PDF del documento. Extraé exactamente: "
        "la fecha de envío (formato DD/MM/AAAA si figura), el número de seguimiento/CD, y el "
        "contenido textual completo del mensaje enviado (no incluyas datos del formulario de "
        "Correo Argentino que no sean parte del texto del mensaje, como códigos de barra). "
        "Transcribí el contenido exactamente como está escrito, sin resumir ni corregir. Si algún "
        "dato no se puede leer con confianza, dejalo vacío en vez de adivinar."
    )

    try:
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-sonnet-5",
            max_tokens=2048,
            system=system_prompt,
            output_config={"format": {
                "type": "json_schema",
                "schema": {
                    "type": "object",
                    "properties": {
                        "fecha": {"type": "string"},
                        "numero": {"type": "string"},
                        "contenido": {"type": "string"}
                    },
                    "required": ["fecha", "numero", "contenido"],
                    "additionalProperties": False
                }
            }},
            messages=[{"role": "user", "content": [
                content_block,
                {"type": "text", "text": "Extraé fecha, número de seguimiento y contenido de este documento."}
            ]}]
        )
        texto_json = "".join([b.text for b in response.content if b.type == "text"])
        datos = json.loads(texto_json)
    except Exception as e:
        print("Error leyendo telegrama:", str(e))
        return error('No se pudo leer el archivo. Probá con otra foto/escaneo, o cargá los datos a mano.', 500)

    resp = jsonify({"status": "ok", "fecha": datos.get('fecha', ''), "numero": datos.get('numero', ''), "contenido": datos.get('contenido', '')})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp, 200

@app.route('/', methods=['GET'])
def health():
    return "Servidor del Bot activo y funcionando correctamente.", 200

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
